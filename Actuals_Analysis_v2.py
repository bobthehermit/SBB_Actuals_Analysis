# Actuals_Analysis.py
# Streamlit app — Cash & Actuals Compliance
# Ad-Hoc Upload Version: User uploads all 3 files per review.
# Improvements: Rollup Logic for Cash, FTE Analysis, Sticky Header, Detailed Findings

import os
import re
import base64
import pickle
from io import BytesIO
from pathlib import Path
from typing import Optional, Dict, Tuple, List
from datetime import datetime

import json
import pandas as pd
import streamlit as st
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ---------- PAGE CONFIG ----------
st.set_page_config(page_title="Actuals Analysis & Compliance", layout="wide")

# ==========================================
# 1. PATH SETUP
# ==========================================
def _script_dir():
    try:
        return Path(__file__).resolve().parent
    except NameError:
        return Path.cwd()

APP_DIR = _script_dir()
LOGO_FILE = APP_DIR / "300 DPI NM PED Logo JPEG.jpg"
CHECKLIST_FILE = APP_DIR / "Actuals_Checklist.csv"

# EXTERNAL LINKS
SEG_EXTERNAL_LINK = "https://webed.ped.state.nm.us/sites/FileTransfer/SitePages/Home.aspx"
OBMS_LINK = "https://obms.ped.state.nm.us/PED_OBMS/OBMSHome"
DFA_DIST_LINK = "https://www.nmdfa.state.nm.us/local-government/budget-finance-bureau/financial-distributions/"

# ---------- THEME / HEADER ----------
LOGO_LEFT_PATH  = str(LOGO_FILE)
LOGO_RIGHT_PATH = "" 
SIDEBAR_LOGO_PATH = LOGO_LEFT_PATH
HEADER_TITLE = "Actuals Analysis & Compliance"
HEADER_SUB   = "Upload reports (Cash, Revenue, Expenditure) to run automated checks."
MAX_LOGO_HEIGHT_PX = 90
SHOW_HEADER_LOGOS = False
LOGO_LEFT_LINK  = "https://web.ped.nm.gov/bureaus/school-budget-bureau/"
LOGO_RIGHT_LINK = "https://web.ped.nm.gov/bureaus/school-budget-bureau/"

# ---------- HELPER FUNCTIONS: UI ----------

def img_to_base64(image_path: str) -> str:
    try:
        if not os.path.exists(image_path):
            return ""
        with open(image_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except Exception:
        return ""

def render_header(title: str, subtitle: str = "", logo_left: Optional[str] = None, logo_right: Optional[str] = None, logo_left_link: str = "", logo_right_link: str = "", show_logos: bool = True):
    if not show_logos:
        st.markdown(f"# {title}")
        if subtitle: st.markdown(f"*{subtitle}*")
        return

    html_parts = ['<div style="display: flex; justify-content: space-between; align-items: center; padding: 1rem 0;">']
    if logo_left and os.path.exists(logo_left):
        b64_left = img_to_base64(logo_left)
        if b64_left:
            logo_html = f'<img src="data:image/jpeg;base64,{b64_left}" style="max-height: {MAX_LOGO_HEIGHT_PX}px;">'
            if logo_left_link: logo_html = f'<a href="{logo_left_link}" target="_blank">{logo_html}</a>'
            html_parts.append(f'<div>{logo_html}</div>')

    html_parts.append(f'<div style="text-align: center; flex-grow: 1;"><h1>{title}</h1><p>{subtitle}</p></div>')
    html_parts.append('</div>')
    st.markdown(''.join(html_parts), unsafe_allow_html=True)

def render_sidebar_logo(logo_path: Optional[str] = None):
    if logo_path and os.path.exists(logo_path):
        try:
            img = Image.open(logo_path)
            st.sidebar.image(img, use_container_width=True)
        except Exception:
            pass

def inject_links(text: str) -> str:
    """Replaces known text patterns with hyperlinks (FTS, OBMS, DFA only)."""
    if not isinstance(text, str): return text
    
    text = text.replace("Navigate to FTS", f"[Navigate to FTS]({SEG_EXTERNAL_LINK})")
    text = text.replace("Navigate to OBMS", f"[Navigate to OBMS]({OBMS_LINK})")
    
    if "44204" in text:
        text += f" [DFA Distributions Info]({DFA_DIST_LINK})"
    
    return text

# ---------- HELPER FUNCTIONS: DATA LOADING ----------

def load_official_checklist() -> List[Dict]:
    if CHECKLIST_FILE.exists():
        try:
            df = pd.read_csv(CHECKLIST_FILE, encoding='utf-8-sig')
            checklist = []
            for idx, row in df.iterrows():
                checklist.append({
                    "step": int(row['step']),
                    "review_area": str(row['Review Area']),
                    "ucoa_line": str(row['UCOA LINE/AREA of Report']).replace('Picture', '').replace('nan', ''),
                    "applies_to": str(row['Districts/ Charters']).replace('nan', ''),
                    "check": str(row['Check ']),
                    "support": str(row['Support']),
                    "method_notes": str(row['Suggested Method/Notes']),
                    "disapprove": str(row['Disapprove?']),
                    "completed": False,
                    "user_notes": ""
                })
            return checklist
        except Exception as e:
            st.error(f"❌ Error loading checklist CSV: {e}")
            return []
    else:
        st.warning(f"⚠️ Checklist not found at: {CHECKLIST_FILE}")
        return []

def load_report_file(file, name: str) -> Optional[pd.DataFrame]:
    try:
        df = pd.read_csv(file) if file.name.endswith('.csv') else pd.read_excel(file)
        return df
    except Exception as e:
        st.error(f"Error loading {name}: {e}")
        return None

def load_cash_from_excel(file) -> Optional[pd.DataFrame]:
    """
    Load cash report from a raw Excel workbook by reading the 'Summary' tab.
    Handles the standard PED cash report format where:
      - Row 0 = headers (Fund, Line 1 ... Line 12)
      - Row 1+ = fund data
      - Last row = GRAND TOTAL
    Also accepts plain CSV files (backward compatible).
    """
    try:
        # If it's a CSV, just load it directly (backward compatible)
        if file.name.lower().endswith('.csv'):
            return pd.read_csv(file)
        
        # Excel file — read the Summary tab
        xls = pd.ExcelFile(file)
        
        if 'Summary' not in xls.sheet_names:
            st.error(f"❌ Could not find 'Summary' tab in {file.name}. Available tabs: {', '.join(xls.sheet_names)}")
            return None
        
        df = pd.read_excel(xls, sheet_name='Summary', header=0)
        
        # Clean up column names — standardize "Line N - long description..." to "Line N"
        clean_cols = {}
        for col in df.columns:
            col_str = str(col).strip()
            if col_str.lower().startswith('line '):
                # Extract "Line N" from "Line N - Some Description..."
                parts = col_str.split(' - ', 1)
                line_label = parts[0].strip()
                clean_cols[col] = line_label
            elif col_str.lower() == 'fund' or col_str.lower().startswith('fund'):
                clean_cols[col] = 'Fund'
            else:
                clean_cols[col] = col_str
        
        df = df.rename(columns=clean_cols)
        
        # Parse entity name from the Excel file name if possible
        # Format: ENTITYNAME-FYxx-Qx-Cash-Report_-xxx-xxx.xlsx
        entity_hint = file.name.split('-')[0] if '-' in file.name else ""
        if entity_hint:
            st.caption(f"📋 Cash report loaded from: **{file.name}** (Summary tab)")
        
        return df
        
    except Exception as e:
        st.error(f"❌ Error reading cash report: {e}")
        return None

def detect_entity_name(rev_df, exp_df) -> str:
    possible = set()
    if rev_df is not None and 'Entity' in rev_df.columns:
        possible.update(rev_df['Entity'].dropna().unique())
    if exp_df is not None and 'Entity' in exp_df.columns:
        possible.update(exp_df['Entity'].dropna().unique())
    return list(possible)[0] if len(possible) == 1 else ""

# ---------- AUTOMATED VALIDATION LOGIC ----------

def clean_currency_series(series: pd.Series) -> pd.Series:
    clean = series.astype(str).str.replace('$', '', regex=False).str.replace(',', '', regex=False).str.replace('(', '-', regex=False).str.replace(')', '', regex=False).str.strip()
    return pd.to_numeric(clean, errors='coerce').fillna(0.0)

def normalize_fund_col(df: pd.DataFrame) -> Tuple[pd.DataFrame, Optional[str]]:
    df.columns = df.columns.str.strip()
    fund_col = next((c for c in df.columns if 'fund' in c.lower() and 'balance' not in c.lower()), None)
    if fund_col:
        def clean_fund_val(val):
            s = str(val).strip()
            if s.endswith('.0'): s = s[:-2]
            if " - " in s: s = s.split(" - ")[0]
            return s.strip()
        df['Fund_Key'] = df[fund_col].apply(clean_fund_val)
        return df, fund_col
    return df, None

def extract_object_code(val):
    """Extract numeric object code from strings like '43101 - State Equalization Guarantee'"""
    s = str(val).strip()
    if " - " in s:
        s = s.split(" - ")[0]
    return s.strip()

def extract_function_code(val):
    """Extract numeric function code from strings like '1000 - Instruction'"""
    s = str(val).strip()
    if " - " in s:
        s = s.split(" - ")[0]
    return s.strip()

def extract_jobclass_code(val):
    """Extract numeric job class code"""
    s = str(val).strip()
    if " - " in s:
        s = s.split(" - ")[0]
    return s.strip()

def find_col(df, keywords):
    candidates = [c for c in df.columns if not any(x in c.lower() for x in ['budget', 'encumb', 'adjust', 'balance'])]
    for k in keywords:
        match = next((c for c in candidates if k in c.lower()), None)
        if match: return match
    return None

def find_col_include(df, keywords):
    """Find column including budget/balance columns"""
    for k in keywords:
        match = next((c for c in df.columns if k in c.lower()), None)
        if match: return match
    return None

def calculate_rollup_sum(df: pd.DataFrame, fund_key: str, amount_col: str) -> float:
    """Calculates sum for a fund. Handles rollups (e.g. 24000 includes 24xxx)."""
    rollup_funds = ['24000', '25000', '26000', '27000', '28000', '29000']
    
    if fund_key in rollup_funds:
        prefix = fund_key[:2]
        return df[df['Fund_Key'].str.startswith(prefix, na=False)][amount_col].sum()
    else:
        return df[df['Fund_Key'] == fund_key][amount_col].sum()

def run_all_validations(cash_df, revenue_df, expenditure_df, entity_name, is_q1, user_inputs):
    """
    Returns:
        results: Dict[step, List[(status, msg)]] - text findings
        table_findings: Dict[step, dict] - structured table data for display
    """
    results = {} 
    table_findings = {}

    def add_finding(step, msg, is_pass=False):
        if step not in results: results[step] = []
        results[step].append(("PASS" if is_pass else "FLAG", msg))

    def add_table(step, df, title=""):
        if not df.empty:
            table_findings[step] = {'title': title, 'data': df}

    def build_detail_table(flagged_df, cols_map, step, title):
        if flagged_df.empty:
            return
        rows = []
        for _, r in flagged_df.iterrows():
            row = {}
            for display_name, source in cols_map.items():
                val = r.get(source, '')
                if isinstance(val, float):
                    row[display_name] = f"${val:,.2f}" if abs(val) >= 0.01 else "$0.00"
                else:
                    row[display_name] = str(val)
            rows.append(row)
        add_table(step, pd.DataFrame(rows), title)

    # --- PREP DATA ---
    c_df = cash_df.copy() if cash_df is not None else None
    
    r_df = None
    if revenue_df is not None:
        r_df = revenue_df.copy()
        if 'Entity' in r_df.columns and entity_name and entity_name in r_df['Entity'].unique():
            r_df = r_df[r_df['Entity'] == entity_name]
        r_df, r_fund_col = normalize_fund_col(r_df)
        
        # Create Object_Key for cleaner matching
        if 'Object' in r_df.columns:
            r_df['Object_Key'] = r_df['Object'].apply(extract_object_code)
        if 'Function' in r_df.columns:
            r_df['Function_Key'] = r_df['Function'].apply(extract_function_code)
            
        for c in r_df.columns:
            if any(x in c.lower() for x in ['amount', 'actual', 'balance', 'budget', 'ytd', 'encumb']):
                r_df[c] = clean_currency_series(r_df[c])

    e_df = None
    if expenditure_df is not None:
        e_df = expenditure_df.copy()
        if 'Entity' in e_df.columns and entity_name and entity_name in e_df['Entity'].unique():
            e_df = e_df[e_df['Entity'] == entity_name]
        e_df, e_fund_col = normalize_fund_col(e_df)
        
        # Create keys for cleaner matching
        if 'Object' in e_df.columns:
            e_df['Object_Key'] = e_df['Object'].apply(extract_object_code)
        if 'Function' in e_df.columns:
            e_df['Function_Key'] = e_df['Function'].apply(extract_function_code)
        if 'JobClass' in e_df.columns:
            e_df['JobClass_Key'] = e_df['JobClass'].apply(extract_jobclass_code)
            
        for c in e_df.columns:
            if any(x in c.lower() for x in ['amount', 'actual', 'balance', 'budget', 'ytd', 'encumb', 'fte']):
                e_df[c] = clean_currency_series(e_df[c])

    if c_df is not None:
        for col in c_df.columns:
            if "Line" in col: c_df[col] = clean_currency_series(c_df[col])
        c_df, c_fund_col = normalize_fund_col(c_df)

    # ==========================================
    # REVENUE REPORT CHECKS
    # ==========================================
    if r_df is not None:
        per_col = find_col(r_df, ['period', 'actuals period'])
        ytd_col = find_col(r_df, ['ytd', 'actuals ytd'])
        bud_col = find_col_include(r_df, ['adjusted budget'])
        enc_col = find_col_include(r_df, ['encumbrance'])
        
        # --- STEP 6: SEG Period Amount (Auto-display) ---
        if 'Object_Key' in r_df.columns and per_col:
            seg_rows = r_df[r_df['Object_Key'] == '43101']
            if not seg_rows.empty:
                fund_11000_seg = seg_rows[seg_rows['Fund_Key'] == '11000']
                if not fund_11000_seg.empty:
                    seg_period = fund_11000_seg[per_col].sum()
                    seg_ytd = fund_11000_seg[ytd_col].sum() if ytd_col else 0
                    add_finding(6, f"SEG (43101) Fund 11000 - Period: ${seg_period:,.2f}, YTD: ${seg_ytd:,.2f}", is_pass=True)
        
        # --- STEP 7: SEG Budgeted Amount (Auto-display) ---
        if 'Object_Key' in r_df.columns and bud_col:
            seg_rows = r_df[r_df['Object_Key'] == '43101']
            if not seg_rows.empty:
                fund_11000_seg = seg_rows[seg_rows['Fund_Key'] == '11000']
                if not fund_11000_seg.empty:
                    seg_budget = fund_11000_seg[bud_col].sum()
                    add_finding(7, f"SEG (43101) Fund 11000 - Adjusted Budget: ${seg_budget:,.2f}", is_pass=True)
        
        # --- STEP 8: Check for negative values in Actuals Period Amount ---
        if per_col:
            negs = r_df[r_df[per_col] < -0.01]
            if not negs.empty:
                add_finding(8, f"🚩 Found {len(negs)} lines with NEGATIVE Period Amount (see table below)")
                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if 'Function_Key' in r_df.columns:
                    cols['Function'] = 'Function_Key'
                cols['Period Amount'] = per_col
                if ytd_col:
                    cols['YTD Amount'] = ytd_col
                build_detail_table(negs, cols, 8, f"Negative Revenue Period Amounts ({len(negs)} lines)")
            else:
                add_finding(8, "No negative values in Actuals Period Amount", is_pass=True)
        
        # --- STEP 9: Check for negative values in Actuals YTD ---
        if ytd_col:
            negs = r_df[r_df[ytd_col] < -0.01]
            if not negs.empty:
                add_finding(9, f"🚩 Found {len(negs)} lines with NEGATIVE YTD Amount (see table below)")
                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if 'Function_Key' in r_df.columns:
                    cols['Function'] = 'Function_Key'
                cols['YTD Amount'] = ytd_col
                if per_col:
                    cols['Period Amount'] = per_col
                build_detail_table(negs, cols, 9, f"Negative Revenue YTD Amounts ({len(negs)} lines)")
            else:
                add_finding(9, "No negative values in Actuals YTD", is_pass=True)
        
        # --- STEP 10: Check for negative values in Revenue Report (Available Balance) ---
        bal_col = find_col_include(r_df, ['available balance'])
        if bal_col:
            negs = r_df[r_df[bal_col] < -0.01]
            if not negs.empty:
                add_finding(10, f"🚩 Found {len(negs)} lines with NEGATIVE Available Balance (see table below)")
                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if bud_col:
                    cols['Adjusted Budget'] = bud_col
                if ytd_col:
                    cols['YTD Amount'] = ytd_col
                cols['Available Balance'] = bal_col
                build_detail_table(negs, cols, 10, f"Negative Revenue Available Balance ({len(negs)} lines)")
            else:
                add_finding(10, "No negative values in Available Balance", is_pass=True)
        
        # --- STEP 11: Check Object 41980 (Vendor Refunds) across all lines ---
        if 'Object_Key' in r_df.columns and ytd_col:
            refunds = r_df[r_df['Object_Key'] == '41980']
            if not refunds.empty:
                total_refund = refunds[ytd_col].sum()
                if abs(total_refund) > 0.01:
                    details = []
                    for _, row in refunds.iterrows():
                        if abs(row[ytd_col]) > 0.01:
                            details.append(f"Fund {row.get('Fund_Key', 'Unknown')}: ${row[ytd_col]:,.2f}")
                    add_finding(11, f"⚠️ Object 41980 (Vendor Refunds) has activity - Total: ${total_refund:,.2f}. Details: {'; '.join(details)}")
                else:
                    add_finding(11, "Object 41980 has no significant activity", is_pass=True)
            else:
                add_finding(11, "Object 41980 not present in revenue report", is_pass=True)
        
        # --- STEP 12: FIXED - Display SEG amounts for Fund 11000 Object 43101 ---
        if 'Object_Key' in r_df.columns and per_col and ytd_col:
            seg_rows = r_df[(r_df['Fund_Key'] == '11000') & (r_df['Object_Key'] == '43101')]
            if not seg_rows.empty:
                seg_period = seg_rows[per_col].sum()
                seg_ytd = seg_rows[ytd_col].sum()
                add_finding(12, f"SEG (Fund 11000, Object 43101):", is_pass=True)
                add_finding(12, f"  • Actuals Period Amount: ${seg_period:,.2f}", is_pass=True)
                add_finding(12, f"  • Actuals YTD Amount: ${seg_ytd:,.2f}", is_pass=True)
            else:
                add_finding(12, "⚠️ SEG (Object 43101) not found in Fund 11000")
        
        # --- STEP 13: SEG YTD Amount for Fund 11000 Object 43101 ---
        if 'Object_Key' in r_df.columns and ytd_col:
            seg_rows = r_df[(r_df['Fund_Key'] == '11000') & (r_df['Object_Key'] == '43101')]
            if not seg_rows.empty:
                seg_ytd = seg_rows[ytd_col].sum()
                add_finding(13, f"SEG (43101) Fund 11000 YTD Amount: ${seg_ytd:,.2f}", is_pass=True)
            else:
                add_finding(13, "⚠️ SEG (Object 43101) not found in Fund 11000")
        
        # --- STEP 14: SEG Budgeted Amount (Adjusted Budget for Fund 11000 Object 43101) ---
        if 'Object_Key' in r_df.columns and bud_col:
            seg_rows = r_df[(r_df['Fund_Key'] == '11000') & (r_df['Object_Key'] == '43101')]
            if not seg_rows.empty:
                seg_budget = seg_rows[bud_col].sum()
                add_finding(14, f"SEG (43101) Fund 11000 Adjusted Budget: ${seg_budget:,.2f}", is_pass=True)
            else:
                add_finding(14, "⚠️ SEG (Object 43101) not found in Fund 11000 for budget check")

        # --- STEPS 15-16: Impact Aid (44103) Ratio Check ---
        if 'Object_Key' in r_df.columns and ytd_col:
            impact_aid_rows = r_df[r_df['Object_Key'] == '44103']
            if not impact_aid_rows.empty:
                total_impact_aid = impact_aid_rows[ytd_col].sum()
                if total_impact_aid > 0:
                    # Build ratio table
                    ratio_data = []
                    for fund in impact_aid_rows['Fund_Key'].unique():
                        fund_amt = impact_aid_rows[impact_aid_rows['Fund_Key'] == fund][ytd_col].sum()
                        fund_budget = impact_aid_rows[impact_aid_rows['Fund_Key'] == fund][bud_col].sum() if bud_col else 0
                        pct = (fund_amt / total_impact_aid) * 100 if total_impact_aid > 0 else 0
                        ratio_data.append({
                            'Fund': fund,
                            'YTD Amount': fund_amt,
                            'Adjusted Budget': fund_budget,
                            'Percentage': pct
                        })
                    
                    ratio_df = pd.DataFrame(ratio_data)
                    ratio_str = " | ".join([f"Fund {r['Fund']}: ${r['YTD Amount']:,.2f} ({r['Percentage']:.1f}%)" for r in ratio_data])
                    
                    # Step 15: Fund 11000 should be 0-25%
                    fund_11000_pct = ratio_df[ratio_df['Fund'] == '11000']['Percentage'].sum() if '11000' in ratio_df['Fund'].values else 0
                    if fund_11000_pct > 25:
                        add_finding(15, f"🚩 Impact Aid (44103) in Fund 11000: {fund_11000_pct:.1f}% (Should be ≤25%). Distribution: {ratio_str}")
                    else:
                        add_finding(15, f"✅ Impact Aid (44103) Fund 11000: {fund_11000_pct:.1f}% (within 0-25% limit). Distribution: {ratio_str}", is_pass=True)
                    
                    # Step 16: Fund 15100 should be 75-100%
                    fund_15100_pct = ratio_df[ratio_df['Fund'] == '15100']['Percentage'].sum() if '15100' in ratio_df['Fund'].values else 0
                    if fund_15100_pct < 75:
                        add_finding(16, f"🚩 Impact Aid (44103) in Fund 15100: {fund_15100_pct:.1f}% (Should be ≥75%)")
                        
                    else:
                        add_finding(16, f"✅ Impact Aid (44103) Fund 15100: {fund_15100_pct:.1f}% (within 75-100% limit)", is_pass=True)
                    
                    # Show detail table for Impact Aid
                    cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                    if per_col: cols['Period Amount'] = per_col
                    if ytd_col: cols['YTD Amount'] = ytd_col
                    if bud_col: cols['Adjusted Budget'] = bud_col
                    build_detail_table(impact_aid_rows, cols, 15, f"Impact Aid (44103) Detail ({len(impact_aid_rows)} lines)")
            else:
                add_finding(15, "Object 44103 (Impact Aid) not present", is_pass=True)
                add_finding(16, "Object 44103 (Impact Aid) not present", is_pass=True)

        # --- STEPS 17-18: Ad Valorem (41110) Ratio Check - Only Fund 11000 and 15200 ---
        if 'Object_Key' in r_df.columns and ytd_col:
            ad_valorem_rows = r_df[r_df['Object_Key'] == '41110']
            if not ad_valorem_rows.empty:
                # Filter to ONLY Fund 11000 and 15200
                relevant_funds = ['11000', '15200']
                av_filtered = ad_valorem_rows[ad_valorem_rows['Fund_Key'].isin(relevant_funds)]
                
                if not av_filtered.empty:
                    total_av = av_filtered[ytd_col].sum()
                    
                    if total_av > 0:
                        fund_11000_amt = av_filtered[av_filtered['Fund_Key'] == '11000'][ytd_col].sum()
                        fund_15200_amt = av_filtered[av_filtered['Fund_Key'] == '15200'][ytd_col].sum()
                        
                        fund_11000_pct = (fund_11000_amt / total_av) * 100 if total_av > 0 else 0
                        fund_15200_pct = (fund_15200_amt / total_av) * 100 if total_av > 0 else 0
                        
                        # Display the breakdown
                        add_finding(17, f"Ad Valorem (41110) - Fund 11000 & 15200 Only:", is_pass=True)
                        add_finding(17, f"  • Fund 11000: ${fund_11000_amt:,.2f} ({fund_11000_pct:.1f}%)", is_pass=True)
                        add_finding(17, f"  • Fund 15200: ${fund_15200_amt:,.2f} ({fund_15200_pct:.1f}%)", is_pass=True)
                        add_finding(17, f"  • Combined Total: ${total_av:,.2f}", is_pass=True)
                        
                        # Step 17: Fund 11000 should be 0-25%
                        if fund_11000_pct > 25:
                            add_finding(17, f"🚩 Fund 11000 is {fund_11000_pct:.1f}% (Should be ≤25%)")
                        else:
                            add_finding(17, f"✅ Fund 11000 ratio OK: {fund_11000_pct:.1f}% (within 0-25%)", is_pass=True)
                        
                        # Step 18: Fund 15200 should be 75-100%
                        if fund_15200_pct < 75:
                            add_finding(18, f"🚩 Fund 15200 is {fund_15200_pct:.1f}% (Should be ≥75%)")
                        else:
                            add_finding(18, f"✅ Fund 15200 ratio OK: {fund_15200_pct:.1f}% (within 75-100%)", is_pass=True)
                    else:
                        add_finding(17, "Ad Valorem (41110) in Funds 11000/15200: $0 YTD", is_pass=True)
                        add_finding(18, "Ad Valorem (41110) in Funds 11000/15200: $0 YTD", is_pass=True)
                else:
                    add_finding(17, "Object 41110 not present in Fund 11000 or 15200", is_pass=True)
                    add_finding(18, "Object 41110 not present in Fund 11000 or 15200", is_pass=True)
                
                # Show detail table for Ad Valorem
                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if per_col: cols['Period Amount'] = per_col
                if ytd_col: cols['YTD Amount'] = ytd_col
                if bud_col: cols['Adjusted Budget'] = bud_col
                build_detail_table(av_filtered, cols, 17, f"Ad Valorem (41110) Funds 11000/15200 Detail ({len(av_filtered)} lines)")

            else:
                add_finding(17, "Object 41110 (Ad Valorem) not present", is_pass=True)
                add_finding(18, "Object 41110 (Ad Valorem) not present", is_pass=True)

        # --- STEPS 19-20: Object 41113 Check ---
        if 'Object_Key' in r_df.columns and ytd_col:
            rows_41113 = r_df[r_df['Object_Key'] == '41113']
            if not rows_41113.empty:
                total = rows_41113[ytd_col].sum()
                ratio_data = []
                for fund in rows_41113['Fund_Key'].unique():
                    fund_amt = rows_41113[rows_41113['Fund_Key'] == fund][ytd_col].sum()
                    pct = (fund_amt / total) * 100 if total > 0 else 0
                    ratio_data.append({'Fund': fund, 'Amount': fund_amt, 'Percentage': pct})
                ratio_str = " | ".join([f"Fund {r['Fund']}: ${r['Amount']:,.2f} ({r['Percentage']:.1f}%)" for r in ratio_data])
                add_finding(19, f"Object 41113 Distribution: {ratio_str}", is_pass=True)
                add_finding(20, f"Object 41113 Total YTD: ${total:,.2f}", is_pass=True)

                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if per_col: cols['Period Amount'] = per_col
                if ytd_col: cols['YTD Amount'] = ytd_col
                if bud_col: cols['Adjusted Budget'] = bud_col
                build_detail_table(rows_41113, cols, 19, f"Object 41113 Detail ({len(rows_41113)} lines)")

            else:
                add_finding(19, "Object 41113 not present", is_pass=True)
                add_finding(20, "Object 41113 not present", is_pass=True)

        # --- STEPS 21-22: Object 41114 Check ---
        if 'Object_Key' in r_df.columns and ytd_col:
            rows_41114 = r_df[r_df['Object_Key'] == '41114']
            if not rows_41114.empty:
                total = rows_41114[ytd_col].sum()
                ratio_data = []
                for fund in rows_41114['Fund_Key'].unique():
                    fund_amt = rows_41114[rows_41114['Fund_Key'] == fund][ytd_col].sum()
                    pct = (fund_amt / total) * 100 if total > 0 else 0
                    ratio_data.append({'Fund': fund, 'Amount': fund_amt, 'Percentage': pct})
                ratio_str = " | ".join([f"Fund {r['Fund']}: ${r['Amount']:,.2f} ({r['Percentage']:.1f}%)" for r in ratio_data])
                add_finding(21, f"Object 41114 Distribution: {ratio_str}", is_pass=True)
                add_finding(22, f"Object 41114 Total YTD: ${total:,.2f}", is_pass=True)

                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if per_col: cols['Period Amount'] = per_col
                if ytd_col: cols['YTD Amount'] = ytd_col
                if bud_col: cols['Adjusted Budget'] = bud_col
                build_detail_table(rows_41114, cols, 21, f"Object 41114 Detail ({len(rows_41114)} lines)")

            else:
                add_finding(21, "Object 41114 not present", is_pass=True)
                add_finding(22, "Object 41114 not present", is_pass=True)

        # --- STEPS 23-24: Forest Reserve (44204) Check ---
        if 'Object_Key' in r_df.columns and ytd_col:
            forest_rows = r_df[r_df['Object_Key'] == '44204']
            if not forest_rows.empty:
                total = forest_rows[ytd_col].sum()
                if total > 0:
                    ratio_data = []
                    for fund in forest_rows['Fund_Key'].unique():
                        fund_amt = forest_rows[forest_rows['Fund_Key'] == fund][ytd_col].sum()
                        fund_budget = forest_rows[forest_rows['Fund_Key'] == fund][bud_col].sum() if bud_col else 0
                        pct = (fund_amt / total) * 100 if total > 0 else 0
                        ratio_data.append({'Fund': fund, 'Amount': fund_amt, 'Budget': fund_budget, 'Percentage': pct})
                    
                    ratio_str = " | ".join([f"Fund {r['Fund']}: ${r['Amount']:,.2f} ({r['Percentage']:.1f}%)" for r in ratio_data])
                    
                    # Step 23: Fund 11000 should be ≤25%
                    fund_11000_pct = next((r['Percentage'] for r in ratio_data if r['Fund'] == '11000'), 0)
                    if fund_11000_pct > 25:
                        add_finding(23, f"🚩 Forest Reserve (44204) in Fund 11000: {fund_11000_pct:.1f}% (Should be ≤25%). Distribution: {ratio_str}")
                    else:
                        add_finding(23, f"✅ Forest Reserve (44204) Fund 11000: {fund_11000_pct:.1f}%. Distribution: {ratio_str}", is_pass=True)
                    
                    # Step 24: Fund 15200 should be ≥75%
                    fund_15200_pct = next((r['Percentage'] for r in ratio_data if r['Fund'] == '15200'), 0)
                    if fund_15200_pct < 75:
                        add_finding(24, f"🚩 Forest Reserve (44204) in Fund 15200: {fund_15200_pct:.1f}% (Should be ≥75%)")
                    else:
                        add_finding(24, f"✅ Forest Reserve (44204) Fund 15200: {fund_15200_pct:.1f}%", is_pass=True)
                    
                    # Show detail table for Forest Reserve
                    cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                    if per_col: cols['Period Amount'] = per_col
                    if ytd_col: cols['YTD Amount'] = ytd_col
                    if bud_col: cols['Adjusted Budget'] = bud_col
                    build_detail_table(forest_rows, cols, 23, f"Forest Reserve (44204) Detail ({len(forest_rows)} lines)")
            else:
                add_finding(23, "Object 44204 (Forest Reserve) not present", is_pass=True)
                add_finding(24, "Object 44204 (Forest Reserve) not present", is_pass=True)

        # --- STEP 25: Fund 21100 YTD Check ---
        if ytd_col:
            fund_21100_rows = r_df[r_df['Fund_Key'] == '21100']
            if not fund_21100_rows.empty:
                total_ytd = fund_21100_rows[ytd_col].sum()
                if abs(total_ytd) > 0.01:
                    add_finding(25, f"✅ Fund 21100 has YTD activity: ${total_ytd:,.2f}", is_pass=True)
                else:
                    add_finding(25, f"🚩 Fund 21100 has $0 YTD - Expected activity for Universal Free Lunch")
                
                cols = {'Fund': 'Fund_Key'}
                if 'Object_Key' in r_df.columns: cols['Object'] = 'Object_Key'
                if per_col: cols['Period Amount'] = per_col
                if ytd_col: cols['YTD Amount'] = ytd_col
                if bud_col: cols['Adjusted Budget'] = bud_col
                build_detail_table(fund_21100_rows, cols, 25, f"Fund 21100 Revenue Detail ({len(fund_21100_rows)} lines)")

            else:
                add_finding(25, "⚠️ Fund 21100 not present in Revenue Report")

    # ==========================================
    # EXPENDITURE REPORT CHECKS
    # ==========================================
    if e_df is not None:
        per_col = find_col(e_df, ['period', 'actuals period'])
        ytd_col = find_col(e_df, ['ytd', 'actuals ytd'])
        enc_col = find_col_include(e_df, ['encumbrance'])
        fte_col = find_col_include(e_df, ['actuals fte'])
        bud_col = find_col_include(e_df, ['adjusted budget'])
        bal_col = find_col_include(e_df, ['available balance'])
        
        # --- STEP 28: Negative Current Period Expenditures --- ALL ROWS ---
        if per_col:
            negs = e_df[e_df[per_col] < -0.01]
            if not negs.empty:
                add_finding(28, f"🚩 Found {len(negs)} lines with NEGATIVE Period Expenditures (see table below)")
                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if 'Function_Key' in e_df.columns:
                    cols['Function'] = 'Function_Key'
                if 'JobClass_Key' in e_df.columns:
                    cols['Job Class'] = 'JobClass_Key'
                cols['Period Amount'] = per_col
                if ytd_col:
                    cols['YTD Amount'] = ytd_col
                build_detail_table(negs, cols, 28, f"Negative Expenditure Period Amounts ({len(negs)} lines)")
            else:
                add_finding(28, "No negative values in Expenditure Period Amount", is_pass=True)
        
        # --- STEP 29: Negative YTD Expenditures --- ALL ROWS ---
        if ytd_col:
            negs = e_df[e_df[ytd_col] < -0.01]
            if not negs.empty:
                add_finding(29, f"🚩 Found {len(negs)} lines with NEGATIVE YTD Expenditures (see table below)")
                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if 'Function_Key' in e_df.columns:
                    cols['Function'] = 'Function_Key'
                if 'JobClass_Key' in e_df.columns:
                    cols['Job Class'] = 'JobClass_Key'
                cols['YTD Amount'] = ytd_col
                if per_col:
                    cols['Period Amount'] = per_col
                build_detail_table(negs, cols, 29, f"Negative Expenditure YTD Amounts ({len(negs)} lines)")
            else:
                add_finding(29, "No negative values in Expenditure YTD", is_pass=True)
        
        # --- STEP 30: Negative Encumbrance Values --- ALL ROWS ---
        if enc_col:
            negs = e_df[e_df[enc_col] < -0.01]
            if not negs.empty:
                add_finding(30, f"🚩 Found {len(negs)} lines with NEGATIVE Encumbrances (see table below)")
                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if 'Function_Key' in e_df.columns:
                    cols['Function'] = 'Function_Key'
                cols['Encumbrance'] = enc_col
                if ytd_col:
                    cols['YTD Amount'] = ytd_col
                build_detail_table(negs, cols, 30, f"Negative Encumbrances ({len(negs)} lines)")
            else:
                add_finding(30, "No negative encumbrance values", is_pass=True)
        
        # --- STEP 31: Object 51100 with $0 YTD and >0 FTE --- ALL ROWS ---
        if 'Object_Key' in e_df.columns and ytd_col and fte_col:
            salary_rows = e_df[e_df['Object_Key'] == '51100']
            bad_rows = salary_rows[(salary_rows[ytd_col] == 0) & (salary_rows[fte_col] > 0)]
            if not bad_rows.empty:
                add_finding(31, f"🚩 Found {len(bad_rows)} salary lines with $0 YTD but >0 FTE (see table below)")
                cols = {'Fund': 'Fund_Key'}
                if 'Function_Key' in e_df.columns:
                    cols['Function'] = 'Function_Key'
                cols['Job Class'] = 'JobClass_Key' if 'JobClass_Key' in e_df.columns else 'Object_Key'
                cols['FTE'] = fte_col
                cols['YTD Amount'] = ytd_col
                build_detail_table(bad_rows, cols, 31, f"Salary Lines with $0 YTD and >0 FTE ({len(bad_rows)} lines)")
            else:
                add_finding(31, "No 51100 lines with $0 YTD and positive FTE", is_pass=True)
        
        # --- STEP 32-33: Object 51100 FTE Validation (excluding substitutes 1610-1613, 1800) ---
        if 'Object_Key' in e_df.columns and 'JobClass_Key' in e_df.columns and ytd_col and fte_col:
            salary_rows = e_df[e_df['Object_Key'] == '51100']
            excluded_jobs = ['1610', '1611', '1612', '1613', '1800']
            
            # Filter to non-excluded job classes
            relevant = salary_rows[~salary_rows['JobClass_Key'].isin(excluded_jobs)]
            
            # Check: YTD > 0 but FTE = 0
            bad_no_fte = relevant[(relevant[ytd_col] > 100) & (relevant[fte_col] == 0)]
            if not bad_no_fte.empty:
                add_finding(32, f"🚩 Found {len(bad_no_fte)} salary lines with >$100 YTD but 0 FTE (see table below)")
                cols = {'Fund': 'Fund_Key'}
                if 'Function_Key' in e_df.columns:
                    cols['Function'] = 'Function_Key'
                cols['Job Class'] = 'JobClass_Key'
                cols['YTD Amount'] = ytd_col
                cols['FTE'] = fte_col
                build_detail_table(bad_no_fte, cols, 32, f"Salary Lines with YTD >$100 but 0 FTE ({len(bad_no_fte)} lines)")
            else:
                add_finding(32, "All salary lines with YTD >$100 have corresponding FTE", is_pass=True)
            
            # --- STEP 32 DETAIL TABLE: Full 51100 by Job Class (excl. subs) ---
            # Statute 22-8-13.3 B, 22-8-48 | PSAB Supplement 14 | SCM Manual | 925F Report
            adj_fte_col_32 = find_col_include(e_df, ['adjusted fte'])
            
            detail_rows_32 = []
            for jc in sorted(relevant['JobClass_Key'].unique()):
                jc_data = relevant[relevant['JobClass_Key'] == jc]
                ytd_sum = jc_data[ytd_col].sum()
                fte_sum = jc_data[fte_col].sum()
                bud_sum = jc_data[bud_col].sum() if bud_col else 0
                enc_sum = jc_data[enc_col].sum() if enc_col else 0
                adj_fte_sum = jc_data[adj_fte_col_32].sum() if adj_fte_col_32 else 0
                
                fte_delta = adj_fte_sum - fte_sum
                avg_salary = ytd_sum / fte_sum if fte_sum > 0 else 0
                
                jc_label = jc
                if 'JobClass' in jc_data.columns:
                    first = jc_data['JobClass'].iloc[0]
                    if ' - ' in str(first):
                        jc_label = str(first)
                
                detail_rows_32.append({
                    'Job Class': jc_label,
                    'Adj Budget': f"${bud_sum:,.2f}",
                    'YTD Salary': f"${ytd_sum:,.2f}",
                    'Encumbered': f"${enc_sum:,.2f}",
                    'Actuals FTE': f"{fte_sum:.2f}",
                    'Budgeted FTE': f"{adj_fte_sum:.2f}",
                    'FTE Δ': f"{fte_delta:+.2f}",
                    'Avg Salary': f"${avg_salary:,.2f}" if fte_sum > 0 else "N/A"
                })
            
            if detail_rows_32:
                add_finding(32, f"📋 Object 51100 by Job Class ({len(detail_rows_32)} classes, excl. substitutes)", is_pass=True)
                add_table(32, pd.DataFrame(detail_rows_32),
                         f"Object 51100 Salary by Job Class ({len(detail_rows_32)} classes, excl. substitutes)")

        # --- STEP 33: FIXED - Check substitute job classes (1610, 1611, 1612, 1613, 1800) for FTE ---
        # These substitute positions should NOT have FTE reported
        if 'Object_Key' in e_df.columns and 'JobClass_Key' in e_df.columns and ytd_col and fte_col:
            salary_rows = e_df[e_df['Object_Key'] == '51100']
            substitute_jobs = ['1610', '1611', '1612', '1613', '1800']
                
            # Filter to ONLY substitute job classes
            substitute_rows = salary_rows[salary_rows['JobClass_Key'].isin(substitute_jobs)]
            
            if not substitute_rows.empty:
                # Check for any FTE > 0 in substitute positions
                bad_fte = substitute_rows[substitute_rows[fte_col] > 0]
                if not bad_fte.empty:
                    add_finding(33, f"🚩 Found {len(bad_fte)} substitute positions with FTE present (see table below)")
                    cols = {'Fund': 'Fund_Key'}
                    if 'Function_Key' in e_df.columns:
                        cols['Function'] = 'Function_Key'
                    cols['Job Class'] = 'JobClass_Key'
                    cols['FTE'] = fte_col
                    cols['YTD Amount'] = ytd_col
                    build_detail_table(bad_fte, cols, 33, f"Substitute Positions with FTE ({len(bad_fte)} lines)")
                else:
                    add_finding(33, f"✅ Substitute positions (Job Classes 1610-1613, 1800) have no FTE as expected ({len(substitute_rows)} lines checked)", is_pass=True)
            else:
                add_finding(33, "No substitute positions (Job Classes 1610-1613, 1800) found in Object 51100", is_pass=True)
        
        # --- STEP 34: Objects 51200 and 51300 must have 0 FTE ---
        if 'Object_Key' in e_df.columns and fte_col:
            non_salary = e_df[e_df['Object_Key'].isin(['51200', '51300'])]
            bad_fte = non_salary[non_salary[fte_col] > 0]
            if not bad_fte.empty:
                add_finding(34, f"🚩 Found {len(bad_fte)} lines with FTE in Objects 51200/51300 (see table below)")
                cols = {'Fund': 'Fund_Key', 'Object': 'Object_Key'}
                if 'Function_Key' in e_df.columns:
                    cols['Function'] = 'Function_Key'
                cols['FTE'] = fte_col
                if ytd_col:
                    cols['YTD Amount'] = ytd_col
                build_detail_table(bad_fte, cols, 34, f"Objects 51200/51300 with FTE ({len(bad_fte)} lines)")
            else:
                add_finding(34, "Objects 51200/51300 have 0 FTE as expected", is_pass=True)
        
        # --- STEP 35: FIXED - Negative Budget Balance at Function SUBTOTAL Level BY FUND ---
        if 'Function_Key' in e_df.columns and bal_col:
            # Group by Fund AND Function to get subtotals
            func_subtotals = e_df.groupby(['Fund_Key', 'Function_Key'])[bal_col].sum().reset_index()
            func_subtotals.columns = ['Fund', 'Function', 'Subtotal_Balance']
            
            # Find fund/function combinations with negative subtotal balance
            neg_subtotals = func_subtotals[func_subtotals['Subtotal_Balance'] < -0.01]
            
            if not neg_subtotals.empty:
                add_finding(35, f"🚩 Found {len(neg_subtotals)} Fund/Function combination(s) with NEGATIVE subtotal balance (see table below)")
                table_rows = []
                for _, row in neg_subtotals.iterrows():
                    table_rows.append({
                        'Fund': str(row['Fund']),
                        'Function': str(row['Function']),
                        'Subtotal Balance': f"${row['Subtotal_Balance']:,.2f}"
                    })
                add_table(35, pd.DataFrame(table_rows), f"Negative Fund/Function Subtotal Balances ({len(neg_subtotals)} combinations)")
            else:
                add_finding(35, "✅ All Fund/Function subtotals have positive or zero Available Balance", is_pass=True)
        
        # --- STEP 36: Object 53330 in Function 1000/2100 (Should be $0) ---
        if 'Object_Key' in e_df.columns and 'Function_Key' in e_df.columns and ytd_col:
            bad_rows = e_df[(e_df['Object_Key'] == '53330') & (e_df['Function_Key'].isin(['1000', '2100']))]
            if not bad_rows.empty:
                total = bad_rows[ytd_col].sum()
                budget = bad_rows[bud_col].sum() if bud_col else 0
                if total > 0 or budget > 0:
                    add_finding(36, f"🚩 Object 53330 in Function 1000/2100: YTD=${total:,.2f}, Budget=${budget:,.2f} (Should be $0)")
                else:
                    add_finding(36, "Object 53330 in Function 1000/2100 has $0 as expected", is_pass=True)
            else:
                add_finding(36, "Object 53330 not used in Functions 1000/2100", is_pass=True)
        
        # --- STEP 37: Forbidden Objects Check (58213, 58214, 58215, 58216, 58218) ---
        # These objects may be BUDGETED but must NOT have actual expenditures (YTD or Period)
        forbidden_objs = ['58213', '58214', '58215', '58216', '58218']
        if 'Object_Key' in e_df.columns and ytd_col:
            bad_rows = e_df[e_df['Object_Key'].isin(forbidden_objs)]
            if not bad_rows.empty:
                expend_flags = []
                budget_only = []
                for obj in forbidden_objs:
                    obj_rows = bad_rows[bad_rows['Object_Key'] == obj]
                    if not obj_rows.empty:
                        ytd_total = obj_rows[ytd_col].sum()
                        per_total = obj_rows[per_col].sum() if per_col else 0
                        bud_total = obj_rows[bud_col].sum() if bud_col else 0
                        has_expenditures = ytd_total > 0.01 or per_total > 0.01
                        
                        if has_expenditures:
                            expend_flags.append(f"Object {obj}: YTD=${ytd_total:,.2f}, Period=${per_total:,.2f}, Budget=${bud_total:,.2f}")
                        elif bud_total > 0.01:
                            budget_only.append(f"Object {obj}: Budget=${bud_total:,.2f} (no expenditures — OK)")
                
                if expend_flags:
                    add_finding(37, f"🚩 Forbidden Objects with EXPENDITURES: {'; '.join(expend_flags)}")
                else:
                    add_finding(37, "✅ No expenditures in forbidden objects (58213-58218)", is_pass=True)
                
                if budget_only:
                    add_finding(37, f"ℹ️ Budgeted only (allowed): {'; '.join(budget_only)}", is_pass=True)
            else:
                add_finding(37, "✅ No forbidden objects (58213-58218) found", is_pass=True)
        
        # --- STEP 38: Fund 11000 Function 4000 Expenditures ---
        if 'Function_Key' in e_df.columns and ytd_col:
            bad_rows = e_df[(e_df['Fund_Key'] == '11000') & (e_df['Function_Key'] == '4000')]
            if not bad_rows.empty:
                total = bad_rows[ytd_col].sum()
                budget = bad_rows[bud_col].sum() if bud_col else 0
                if total > 0:
                    add_finding(38, f"🚩 Fund 11000 Function 4000 has expenditures: YTD=${total:,.2f}, Budget=${budget:,.2f}")
                else:
                    add_finding(38, f"Fund 11000 Function 4000: YTD=$0, Budget=${budget:,.2f}", is_pass=True)
            else:
                add_finding(38, "Fund 11000 Function 4000 not present", is_pass=True)
        
        # --- STEP 39: Fund 21100 Expenditures (Should have activity) ---
        if ytd_col:
            fund_21100 = e_df[e_df['Fund_Key'] == '21100']
            if not fund_21100.empty:
                total = fund_21100[ytd_col].sum()
                if total > 0:
                    add_finding(39, f"✅ Fund 21100 has expenditures: ${total:,.2f}", is_pass=True)
                else:
                    add_finding(39, f"🚩 Fund 21100 has $0 YTD Expenditures (Expected activity)")
            else:
                add_finding(39, "⚠️ Fund 21100 not present in Expenditure Report")
        
        # --- STEP 40: Q1 Period vs YTD Check (Expenditure) ---
        if per_col and ytd_col:
            grand_total_period = e_df[per_col].sum()
            grand_total_ytd = e_df[ytd_col].sum()
            
            add_finding(40, f"Expenditure Grand Totals - Period: ${grand_total_period:,.2f}, YTD: ${grand_total_ytd:,.2f}", is_pass=True)
            
            if is_q1:
                if abs(grand_total_period - grand_total_ytd) < 1.0:
                    add_finding(40, "✅ Q1 Check PASSED: Period equals YTD", is_pass=True)
                else:
                    add_finding(40, f"🚩 Q1 Check: Period (${grand_total_period:,.2f}) != YTD (${grand_total_ytd:,.2f})")
            else:
                if abs(grand_total_period - grand_total_ytd) < 1.0:
                    add_finding(40, f"🚩 Q2-Q4 Flag: Period equals YTD (unexpected)")
                else:
                    add_finding(40, "✅ Q2-Q4: Period and YTD are different (expected)", is_pass=True)

    # ==========================================
    # CASH REPORT CHECKS
    # ==========================================
    if c_df is not None:
        # --- STEP 48: Revenue (Line 2) vs Revenue Report (Rollup Logic) - WITH DETAILED MATH ---
        col_l2 = next((c for c in c_df.columns if 'Line 2' in c), None)
        if col_l2 and r_df is not None:
            r_amt = find_col(r_df, ['ytd', 'actuals ytd'])
            if r_amt:
                # Grand Total Check
                cash_grand_total_rev = c_df[c_df['Fund_Key'] != 'GRAND TOTAL'][col_l2].sum()
                rev_report_grand_total = r_df[r_amt].sum()
                
                add_finding(48, f"📊 Revenue Reconciliation (Cash Line 2 vs Revenue Report YTD):", is_pass=True)
                add_finding(48, f"  Cash Report Grand Total: ${cash_grand_total_rev:,.2f}", is_pass=True)
                add_finding(48, f"  Revenue Report Grand Total: ${rev_report_grand_total:,.2f}", is_pass=True)
                
                diff = cash_grand_total_rev - rev_report_grand_total
                if abs(diff) > 1.0:
                    add_finding(48, f"🚩 GRAND TOTAL DIFFERENCE: ${diff:,.2f}")
                else:
                    add_finding(48, f"✅ Grand Totals Match (Difference: ${diff:,.2f})", is_pass=True)
                
                # Per-fund detailed breakdown
                add_finding(48, f"", is_pass=True)
                add_finding(48, f"📋 Per-Fund Comparison (Cash Line 2 vs Revenue Report YTD):", is_pass=True)
                
                matches = []
                mismatches = []
                
                for _, row in c_df.iterrows():
                    fund = str(row['Fund_Key'])
                    if fund == 'GRAND TOTAL':
                        continue
                    
                    cash_val = row[col_l2]
                    
                    # Use rollup logic for 24xxx, 25xxx, 26xxx, 27xxx, 28xxx, 29xxx funds
                    rev_val = calculate_rollup_sum(r_df, fund, r_amt)
                    
                    diff = cash_val - rev_val
                    
                    if abs(diff) > 1.0:
                        mismatches.append({
                            'fund': fund,
                            'cash': cash_val,
                            'rev': rev_val,
                            'diff': diff
                        })
                    else:
                        matches.append({
                            'fund': fund,
                            'cash': cash_val,
                            'rev': rev_val
                        })
                
                # Show mismatches first (these need attention)
                if mismatches:
                    add_finding(48, f"🚩 Funds with DIFFERENCES ({len(mismatches)}) - Needs Review:")
                    for m in mismatches:
                        add_finding(48, f"  • Fund {m['fund']}:")
                        add_finding(48, f"      Cash Line 2: ${m['cash']:,.2f}")
                        add_finding(48, f"      Revenue Rpt: ${m['rev']:,.2f}")
                        add_finding(48, f"      Difference:  ${m['diff']:,.2f}")
                
                # Show matching funds
                if matches:
                    add_finding(48, f"✅ Funds that MATCH ({len(matches)}):", is_pass=True)
                    for m in matches:
                        # Only show detail if there's actual activity
                        if m['cash'] != 0 or m['rev'] != 0:
                            add_finding(48, f"  • Fund {m['fund']}: Cash=${m['cash']:,.2f} = Rev=${m['rev']:,.2f}", is_pass=True)
                        else:
                            add_finding(48, f"  • Fund {m['fund']}: $0.00 (no activity)", is_pass=True)

        # --- STEP 49: Expenditure (Line 5) vs Expenditure Report (Rollup Logic) - WITH DETAILED MATH ---
        col_l5 = next((c for c in c_df.columns if 'Line 5' in c), None)
        if col_l5 and e_df is not None:
            e_amt = find_col(e_df, ['ytd', 'actuals ytd'])
            if e_amt:
                # Note: Line 5 in Cash is NEGATIVE, Expenditure Report YTD is POSITIVE
                cash_grand_total_exp = c_df[c_df['Fund_Key'] != 'GRAND TOTAL'][col_l5].sum()  # Negative
                exp_report_grand_total = e_df[e_amt].sum()  # Positive
                
                add_finding(49, f"📊 Expenditure Reconciliation (|Cash Line 5| vs Expenditure Report YTD):", is_pass=True)
                add_finding(49, f"  Cash Report |Line 5|: ${abs(cash_grand_total_exp):,.2f}", is_pass=True)
                add_finding(49, f"  Expenditure Report Grand Total: ${exp_report_grand_total:,.2f}", is_pass=True)
                
                diff = abs(cash_grand_total_exp) - exp_report_grand_total
                if abs(diff) > 1.0:
                    add_finding(49, f"🚩 GRAND TOTAL DIFFERENCE: ${diff:,.2f}")
                else:
                    add_finding(49, f"✅ Grand Totals Match (Difference: ${diff:,.2f})", is_pass=True)
                
                # Per-fund detailed breakdown
                add_finding(49, f"", is_pass=True)
                add_finding(49, f"📋 Per-Fund Comparison (|Cash Line 5| vs Expenditure Report YTD):", is_pass=True)
                
                matches = []
                mismatches = []
                
                for _, row in c_df.iterrows():
                    fund = str(row['Fund_Key'])
                    if fund == 'GRAND TOTAL':
                        continue
                    
                    cash_val = abs(row[col_l5])  # Convert to positive for comparison
                    
                    # Use rollup logic for 24xxx, 25xxx, 26xxx, 27xxx, 28xxx, 29xxx funds
                    exp_val = calculate_rollup_sum(e_df, fund, e_amt)
                    
                    diff = cash_val - exp_val
                    
                    if abs(diff) > 1.0:
                        mismatches.append({
                            'fund': fund,
                            'cash': cash_val,
                            'exp': exp_val,
                            'diff': diff
                        })
                    else:
                        matches.append({
                            'fund': fund,
                            'cash': cash_val,
                            'exp': exp_val
                        })
                
                # Show mismatches first (these need attention)
                if mismatches:
                    add_finding(49, f"🚩 Funds with DIFFERENCES ({len(mismatches)}) - Needs Review:")
                    for m in mismatches:
                        add_finding(49, f"  • Fund {m['fund']}:")
                        add_finding(49, f"      |Cash Line 5|: ${m['cash']:,.2f}")
                        add_finding(49, f"      Expenditure Rpt: ${m['exp']:,.2f}")
                        add_finding(49, f"      Difference:  ${m['diff']:,.2f}")
                
                # Show matching funds
                if matches:
                    add_finding(49, f"✅ Funds that MATCH ({len(matches)}):", is_pass=True)
                    for m in matches:
                        # Only show detail if there's actual activity
                        if m['cash'] != 0 or m['exp'] != 0:
                            add_finding(49, f"  • Fund {m['fund']}: |Cash|=${m['cash']:,.2f} = Exp=${m['exp']:,.2f}", is_pass=True)
                        else:
                            add_finding(49, f"  • Fund {m['fund']}: $0.00 (no activity)", is_pass=True)

        # Step 50: Revenue vs Expenditure by Fund
        if col_l2 and col_l5:
            for _, row in c_df.iterrows():
                fund = str(row['Fund_Key'])
                if fund == 'GRAND TOTAL':
                    continue
                net = row[col_l2] + row[col_l5]  # Rev (pos) + Exp (neg)
                if net < -0.01:
                    add_finding(50, f"🚩 Fund {fund}: Expenditures exceed Revenue (Net: ${net:,.2f})")

        # Step 51: Cash Transfers (Line 6)
        col_l6 = next((c for c in c_df.columns if 'Line 6' in c), None)
        if col_l6:
            transfers = []
            for _, row in c_df.iterrows():
                if abs(row[col_l6]) > 0.01 and row['Fund_Key'] != 'GRAND TOTAL':
                    transfers.append(f"Fund {row['Fund_Key']}: ${row[col_l6]:,.2f}")
            if transfers:
                add_finding(51, f"⚠️ Cash Transfers on Line 6: {'; '.join(transfers)} - Verify with OBMS/FTS")
            else:
                add_finding(51, "No cash transfers on Line 6", is_pass=True)

        # Step 52: Negative Cash Balance (Line 7) --- TABLE FORMAT ---
        col_l7 = next((c for c in c_df.columns if 'Line 7' in c), None)
        if col_l7:
            neg_cash_rows = []
            for _, row in c_df.iterrows():
                if row[col_l7] < -0.01 and row['Fund_Key'] != 'GRAND TOTAL':
                    neg_cash_rows.append({
                        'Fund': str(row['Fund_Key']),
                        'Cash Balance (Line 7)': f"${row[col_l7]:,.2f}"
                    })
            if neg_cash_rows:
                add_finding(52, f"🚩 Negative Cash Balance (Line 7) in {len(neg_cash_rows)} fund(s) (see table below)")
                add_table(52, pd.DataFrame(neg_cash_rows), f"Negative Cash Balances ({len(neg_cash_rows)} funds)")
            else:
                add_finding(52, "No negative cash balances on Line 7", is_pass=True)

        # --- STEP 53: FIXED - Payroll Liability Check (Line 8) ---
        # Find column that contains "Line 8" - handle various naming conventions
        col_l8 = None
        for c in c_df.columns:
            if 'line 8' in c.lower() or ('8' in c and 'payroll' in c.lower()) or ('8' in c and 'liabilit' in c.lower()):
                col_l8 = c
                break
        
        if col_l8 is None:
            # Try more flexible matching
            col_l8 = next((c for c in c_df.columns if 'Line 8' in c), None)
        
        if col_l8:
            add_finding(53, f"Payroll Liabilities (Line 8) Analysis:", is_pass=True)
            
            issues_found = []
            for _, row in c_df.iterrows():
                fund = str(row.get('Fund_Key', ''))
                if fund == 'GRAND TOTAL' or not fund:
                    continue
                
                liability_val = row[col_l8]
                
                # Check for negative values (concerning - should generally be positive)
                if liability_val < -0.01:
                    issues_found.append(f"Fund {fund}: NEGATIVE payroll liability ${liability_val:,.2f}")
                # Note positive values (expected)
                elif liability_val > 0.01:
                    add_finding(53, f"  • Fund {fund}: Payroll liability = ${liability_val:,.2f}", is_pass=True)
            
            if issues_found:
                add_finding(53, "⚠️ Payroll Liability Concerns Found:")
                for issue in issues_found:
                    add_finding(53, f"  🚩 {issue}")
                add_finding(53, "  ℹ️ Note: This is a self-reported cell. Use professional judgment to determine if these are true payroll liabilities or a plug to make the cash report balance. These amounts should generally be positive.")
            else:
                # Check if there are ANY payroll liabilities at all
                total_liabilities = c_df[c_df['Fund_Key'] != 'GRAND TOTAL'][col_l8].sum()
                if abs(total_liabilities) < 0.01:
                    add_finding(53, "ℹ️ No payroll liabilities reported across any fund. Verify this is expected.", is_pass=True)
                else:
                    add_finding(53, f"✅ All payroll liabilities are positive (Total: ${total_liabilities:,.2f})", is_pass=True)
        else:
            add_finding(53, "⚠️ Could not find Line 8 (Payroll Liabilities) column in Cash Report. Available columns: " + ", ".join(c_df.columns.tolist()))

        # Step 54: Interfund Loans (Line 11)
        col_l11 = next((c for c in c_df.columns if 'Line 11' in c), None)
        if col_l11:
            total_loans = c_df[c_df['Fund_Key'] != 'GRAND TOTAL'][col_l11].sum()
            if abs(total_loans) > 0.01:
                add_finding(54, f"🚩 Line 11 Grand Total != $0 (${total_loans:,.2f}) - Interfund loans must balance")
            else:
                add_finding(54, "Interfund loans balance (Line 11 total = $0)", is_pass=True)
            
            # Flag problematic entries
            for _, row in c_df.iterrows():
                fund = str(row['Fund_Key'])
                val = row[col_l11]
                if fund == 'GRAND TOTAL':
                    continue
                if fund == '11000' and val > 0.01:
                    add_finding(54, f"🚩 Fund 11000 has POSITIVE loan (${val:,.2f}) - Should be negative or zero")
                elif fund != '11000' and val < -0.01:
                    add_finding(54, f"🚩 Fund {fund} has NEGATIVE loan (${val:,.2f}) - Should be positive or zero")

        # Step 55: Capital Fund Revenue (31100 and 31900)
        if r_df is not None:
            r_amt = find_col(r_df, ['ytd', 'actuals ytd'])
            if r_amt:
                fund_31100 = r_df[r_df['Fund_Key'] == '31100']
                rev_31100 = fund_31100[r_amt].sum() if not fund_31100.empty else 0
                
                fund_31900 = r_df[r_df['Fund_Key'] == '31900']
                rev_31900 = fund_31900[r_amt].sum() if not fund_31900.empty else 0
                
                add_finding(55, f"Fund 31100 (GO Bond Building) Revenue: ${rev_31100:,.2f}", is_pass=True)
                add_finding(55, f"Fund 31900 Revenue: ${rev_31900:,.2f}", is_pass=True)

    # Step 47: Cash Balance Year-over-Year
    ly = user_inputs.get('step_47_last_year', 0.0)
    ty = user_inputs.get('step_47_this_year', 0.0)
    if ly != 0 and ty != 0 and abs(ly - ty) > 0.01:
        add_finding(47, f"Cash Balance Difference: Last Year ${ly:,.2f} vs This Year ${ty:,.2f}")

    return results, table_findings

# ---------- ANALYSIS SUMMARY GENERATOR ----------

def generate_analysis_summary(cash_df, revenue_df, expenditure_df, entity_name, is_q1, validation_results) -> Dict:
    """Generate comprehensive analysis summary for memo export."""
    summary = {
        'entity': entity_name,
        'review_date': datetime.now().strftime("%B %d, %Y"),
        'period': 'Q1' if is_q1 else 'Q2-Q4',
        'highlights': [],
        'concerns': [],
        'statistics': {}
    }
    
    # Collect statistics
    if revenue_df is not None:
        r_df = revenue_df.copy()
        r_df, _ = normalize_fund_col(r_df)
        ytd_col = find_col(r_df, ['ytd', 'actuals ytd'])
        if ytd_col:
            for c in r_df.columns:
                if 'ytd' in c.lower() or 'amount' in c.lower():
                    r_df[c] = clean_currency_series(r_df[c])
            summary['statistics']['total_revenue'] = r_df[ytd_col].sum()
    
    if expenditure_df is not None:
        e_df = expenditure_df.copy()
        e_df, _ = normalize_fund_col(e_df)
        ytd_col = find_col(e_df, ['ytd', 'actuals ytd'])
        if ytd_col:
            for c in e_df.columns:
                if 'ytd' in c.lower() or 'amount' in c.lower():
                    e_df[c] = clean_currency_series(e_df[c])
            summary['statistics']['total_expenditure'] = e_df[ytd_col].sum()
    
    if cash_df is not None:
        c_df = cash_df.copy()
        c_df, _ = normalize_fund_col(c_df)
        for col in c_df.columns:
            if "Line" in col:
                c_df[col] = clean_currency_series(c_df[col])
        
        col_l7 = next((c for c in c_df.columns if 'Line 7' in c), None)
        if col_l7:
            summary['statistics']['total_cash'] = c_df[c_df['Fund_Key'] != 'GRAND TOTAL'][col_l7].sum()
    
    # Categorize findings
    for step, findings in validation_results.items():
        for finding in findings:
            status, msg = finding
            if status == "FLAG" and "🚩" in msg:
                summary['concerns'].append(f"Step {step}: {msg}")
            elif status == "PASS" and "✅" in msg:
                summary['highlights'].append(f"Step {step}: {msg}")
    
    return summary


# --- HTML VISUAL REPORT GENERATOR ---

# --- UPGRADED HTML VISUAL REPORT GENERATOR (v2) ---
# Replace the existing generate_html_report function in Actuals_Analysis.py with this.
# Same function signature — drop-in replacement.

def generate_html_report(
    entity_name: str,
    revenue_df: pd.DataFrame,
    expenditure_df: pd.DataFrame,
    cash_df: pd.DataFrame,
    validation_results: Dict,
    table_findings: Dict = None,
    notes_by_step: Dict = None,
    checklist_data: List[Dict] = None
) -> str:
    """
    Generates a professional HTML financial analysis report (v2).
    New sections: Encumbrance Risk, Burn Rate, FTE Variance, Program Spending,
    Reimbursable/Non-Reimbursable split, Action Items with due dates.
    """
    
    review_date = datetime.now().strftime("%B %d, %Y")
    now = datetime.now()
    fiscal_year_num = now.year if now.month >= 7 else now.year - 1
    fy_short = str(fiscal_year_num + 1)[2:]  # e.g. "26" for FY26
    fiscal_year = f"FY{fy_short}"
    
    # Determine quarter from review context
    quarter_label = "Q1"  # Default; could be passed as param in future
    
    # Due date calculation
    from datetime import timedelta
    review_dt = now
    high_due = (review_dt + timedelta(days=14)).strftime("%B %d, %Y")  # 10 business ≈ 14 calendar
    med_due = (review_dt + timedelta(days=21)).strftime("%B %d, %Y")   # 15 business ≈ 21 calendar
    low_due = (review_dt + timedelta(days=28)).strftime("%B %d, %Y")   # 20 business ≈ 28 calendar

    # ======================================================
    # DATA PREPARATION
    # ======================================================
    total_revenue = 0
    total_expenditure = 0
    total_encumbrance = 0
    total_budget = 0
    total_cash = 0
    
    # --- Revenue ---
    revenue_by_source = {}
    revenue_by_fund = {}
    r_df = None
    if revenue_df is not None:
        r_df = revenue_df.copy()
        r_df, _ = normalize_fund_col(r_df)
        if 'Object' in r_df.columns:
            r_df['Object_Key'] = r_df['Object'].apply(extract_object_code)
        ytd_col = find_col(r_df, ['ytd', 'actuals ytd'])
        bud_col_r = find_col_include(r_df, ['adjusted budget'])
        if ytd_col:
            for c in r_df.columns:
                if any(x in c.lower() for x in ['amount', 'actual', 'ytd', 'budget']):
                    r_df[c] = clean_currency_series(r_df[c])
            
            # Exclude cash objects from revenue total
            cash_objects = ['11111', '11112']
            rev_mask = ~r_df['Object_Key'].isin(cash_objects) if 'Object_Key' in r_df.columns else pd.Series([True]*len(r_df))
            total_revenue = r_df.loc[rev_mask, ytd_col].sum()
            
            # Revenue by fund (excluding cash)
            for fund in r_df.loc[rev_mask, 'Fund_Key'].unique():
                fund_rev = r_df.loc[rev_mask & (r_df['Fund_Key'] == fund), ytd_col].sum()
                if abs(fund_rev) > 0.01:
                    fund_name = fund
                    if 'Fund' in r_df.columns:
                        first = r_df.loc[r_df['Fund_Key'] == fund, 'Fund'].iloc[0]
                        if ' - ' in str(first):
                            fund_name = str(first).split(' - ', 1)[1][:40]
                    revenue_by_fund[fund] = {'name': fund_name, 'amount': fund_rev}
            
            # Revenue by source type
            if 'Object_Key' in r_df.columns:
                seg = r_df.loc[rev_mask & (r_df['Object_Key'] == '43101'), ytd_col].sum()
                transport = r_df.loc[rev_mask & r_df['Object_Key'].isin(['43206', '43202']), ytd_col].sum()
                federal = r_df.loc[rev_mask & r_df['Object_Key'].str.startswith('44', na=False), ytd_col].sum()
                state_other = r_df.loc[rev_mask & r_df['Object_Key'].str.startswith('43', na=False) & ~r_df['Object_Key'].isin(['43101', '43206', '43202']), ytd_col].sum()
                local = r_df.loc[rev_mask & r_df['Object_Key'].str.startswith('41', na=False), ytd_col].sum()
                revenue_by_source = {
                    'SEG (43101)': seg,
                    'Transportation': transport,
                    'Federal Grants': federal,
                    'Other State': state_other,
                    'Local Revenue': local
                }
    
    # --- Expenditure ---
    expenditure_by_function = {}
    expenditure_by_program = {}
    salary_by_jobclass = {}
    encumbrance_risk_lines = []
    burn_over_pace = []
    burn_under_pace = []
    fte_variances = []
    exp_by_fund = {}
    
    e_df = None
    if expenditure_df is not None:
        e_df = expenditure_df.copy()
        e_df, _ = normalize_fund_col(e_df)
        if 'Object' in e_df.columns:
            e_df['Object_Key'] = e_df['Object'].apply(extract_object_code)
        if 'Function' in e_df.columns:
            e_df['Function_Key'] = e_df['Function'].apply(extract_function_code)
        if 'JobClass' in e_df.columns:
            e_df['JobClass_Key'] = e_df['JobClass'].apply(extract_jobclass_code)
        if 'Program' in e_df.columns:
            e_df['Program_Key'] = e_df['Program'].apply(lambda x: str(x).split(' - ')[0].strip() if pd.notna(x) else '0000')
        
        ytd_col = find_col(e_df, ['ytd', 'actuals ytd'])
        bud_col = find_col_include(e_df, ['adjusted budget'])
        enc_col = find_col_include(e_df, ['encumbrance'])
        fte_col = find_col_include(e_df, ['actuals fte'])
        adj_fte_col = find_col_include(e_df, ['adjusted fte'])
        bal_col = find_col_include(e_df, ['available balance'])
        
        if ytd_col:
            for c in e_df.columns:
                if any(x in c.lower() for x in ['amount', 'actual', 'ytd', 'budget', 'fte', 'encumb', 'balance']):
                    e_df[c] = clean_currency_series(e_df[c])
            
            total_expenditure = e_df[ytd_col].sum()
            total_encumbrance = e_df[enc_col].sum() if enc_col else 0
            total_budget = e_df[bud_col].sum() if bud_col else 0
            
            # --- Expenditure by fund ---
            for fund in e_df['Fund_Key'].unique():
                fund_data = e_df[e_df['Fund_Key'] == fund]
                fund_ytd = fund_data[ytd_col].sum()
                fund_bud = fund_data[bud_col].sum() if bud_col else 0
                fund_enc = fund_data[enc_col].sum() if enc_col else 0
                if fund_ytd > 0 or fund_bud > 0:
                    fund_name = fund
                    if 'Fund' in fund_data.columns:
                        first = fund_data['Fund'].iloc[0]
                        if ' - ' in str(first):
                            fund_name = str(first).split(' - ', 1)[1][:40]
                    exp_by_fund[fund] = {
                        'name': fund_name, 'ytd': fund_ytd,
                        'budget': fund_bud, 'encumbrance': fund_enc
                    }
            
            # --- Expenditure by function (with encumbrance) ---
            func_groups = {
                '1000': 'Instruction', '2100': 'Support Services - Students',
                '2200': 'Support Services - Instruction', '2300': 'General Administration',
                '2400': 'School Administration', '2500': 'Central Services',
                '2600': 'Operation & Maintenance', '2700': 'Student Transportation',
                '2900': 'Other Support Services', '3100': 'Food Services',
                '3300': 'Community Services', '4000': 'Capital Outlay', '5000': 'Debt Service'
            }
            if 'Function_Key' in e_df.columns:
                for func_code, func_name in func_groups.items():
                    func_data = e_df[e_df['Function_Key'] == func_code]
                    if not func_data.empty:
                        ytd_sum = func_data[ytd_col].sum()
                        bud_sum = func_data[bud_col].sum() if bud_col else 0
                        enc_sum = func_data[enc_col].sum() if enc_col else 0
                        avail = bud_sum - ytd_sum - enc_sum
                        pct_used = ((ytd_sum + enc_sum) / bud_sum * 100) if bud_sum > 0 else 0
                        if ytd_sum > 0 or bud_sum > 0:
                            expenditure_by_function[func_code] = {
                                'name': func_name, 'ytd': ytd_sum, 'budget': bud_sum,
                                'encumbrance': enc_sum, 'available': avail, 'pct_used': pct_used
                            }
            
            # --- Expenditure by program ---
            if 'Program_Key' in e_df.columns and 'Program' in e_df.columns:
                program_names = {
                    '0000': 'No Program (Support/Admin)',
                    '1010': 'Regular Education (PreK-12)',
                    '1020': 'Elementary Fine Arts',
                    '2000': 'Special Programs',
                    '3000': 'Vocational & Technical',
                    '9000': 'Co-Curricular & Extra-Curricular'
                }
                for prog in e_df['Program_Key'].unique():
                    prog_data = e_df[e_df['Program_Key'] == prog]
                    ytd_sum = prog_data[ytd_col].sum()
                    bud_sum = prog_data[bud_col].sum() if bud_col else 0
                    enc_sum = prog_data[enc_col].sum() if enc_col else 0
                    if ytd_sum > 0 or bud_sum > 0:
                        prog_name = program_names.get(prog, '')
                        if not prog_name and 'Program' in prog_data.columns:
                            first = prog_data['Program'].iloc[0]
                            prog_name = str(first).split(' - ', 1)[1][:40] if ' - ' in str(first) else prog
                        expenditure_by_program[prog] = {
                            'name': prog_name, 'ytd': ytd_sum,
                            'budget': bud_sum, 'encumbrance': enc_sum,
                            'pct_used': ((ytd_sum + enc_sum) / bud_sum * 100) if bud_sum > 0 else 0
                        }
            
            # --- Encumbrance Risk Analysis ---
            # Lines where Actuals + Encumbrances > Budget
            if enc_col and bud_col:
                e_df['_total_committed'] = e_df[ytd_col] + e_df[enc_col]
                e_df['_over_committed'] = e_df['_total_committed'] - e_df[bud_col]
                over_committed = e_df[e_df['_over_committed'] > 100].copy()  # threshold $100
                over_committed = over_committed.sort_values('_over_committed', ascending=False).head(20)
                
                for _, row in over_committed.iterrows():
                    fund_label = row.get('Fund_Key', '')
                    func_label = ''
                    if 'Function' in row.index:
                        func_label = str(row['Function']).split(' - ', 1)[1][:25] if ' - ' in str(row.get('Function', '')) else row.get('Function_Key', '')
                    obj_label = ''
                    if 'Object' in row.index:
                        obj_label = str(row['Object']).split(' - ', 1)[1][:30] if ' - ' in str(row.get('Object', '')) else row.get('Object_Key', '')
                    
                    encumbrance_risk_lines.append({
                        'fund': fund_label,
                        'function': func_label,
                        'object': obj_label,
                        'budget': row[bud_col],
                        'ytd': row[ytd_col],
                        'encumbered': row[enc_col],
                        'total_committed': row['_total_committed'],
                        'over_amount': row['_over_committed']
                    })
            
            # --- Burn Rate Analysis ---
            if bud_col:
                for _, row in e_df.iterrows():
                    budget = row[bud_col]
                    ytd = row[ytd_col]
                    obj_key = row.get('Object_Key', '')
                    
                    # Skip cash/reserve objects
                    if obj_key in ('11111', '11112', '58214', '58215', '58221'):
                        continue
                    
                    if budget > 0:
                        burn_pct = (ytd / budget) * 100
                        
                        fund_label = row.get('Fund_Key', '')
                        func_label = row.get('Function_Key', '')
                        obj_label = ''
                        if 'Object' in row.index:
                            obj_label = str(row['Object']).split(' - ', 1)[1][:30] if ' - ' in str(row.get('Object', '')) else obj_key
                        
                        line_info = {
                            'fund': fund_label,
                            'function': func_label,
                            'object': obj_label,
                            'budget': budget,
                            'ytd': ytd,
                            'burn_pct': burn_pct
                        }
                        
                        # Over-pace: >40% burn at Q1
                        if burn_pct > 40 and ytd > 1000:
                            burn_over_pace.append(line_info)
                        
                        # Under-pace: 0% burn with budget > $50k
                        if abs(ytd) < 0.01 and budget > 50000:
                            burn_under_pace.append(line_info)
                
                burn_over_pace.sort(key=lambda x: x['burn_pct'], reverse=True)
                burn_over_pace = burn_over_pace[:15]
                burn_under_pace.sort(key=lambda x: x['budget'], reverse=True)
                burn_under_pace = burn_under_pace[:15]
            
            # --- FTE Variance Analysis ---
            if 'Object_Key' in e_df.columns and fte_col and adj_fte_col:
                salary_rows = e_df[e_df['Object_Key'] == '51100'].copy()
                # Exclude substitutes
                if 'JobClass_Key' in salary_rows.columns:
                    salary_rows = salary_rows[~salary_rows['JobClass_Key'].isin(['1610', '1611', '1612', '1613', '1800'])]
                
                # Group by JobClass
                if 'JobClass_Key' in salary_rows.columns:
                    for jc in salary_rows['JobClass_Key'].unique():
                        jc_data = salary_rows[salary_rows['JobClass_Key'] == jc]
                        actual_fte = jc_data[fte_col].sum()
                        budgeted_fte = jc_data[adj_fte_col].sum()
                        variance = actual_fte - budgeted_fte
                        
                        if abs(variance) >= 2.0:
                            jc_name = jc
                            if 'JobClass' in jc_data.columns:
                                first = jc_data['JobClass'].iloc[0]
                                if ' - ' in str(first):
                                    jc_name = str(first).split(' - ', 1)[1][:35]
                            
                            status = 'understaffed' if variance < 0 else 'overstaffed'
                            note = 'Potential vacancies' if variance < 0 else 'May indicate unauthorized positions'
                            
                            fte_variances.append({
                                'jobclass_code': jc,
                                'jobclass': jc_name,
                                'actual_fte': actual_fte,
                                'budgeted_fte': budgeted_fte,
                                'variance': variance,
                                'status': status,
                                'note': note
                            })
                    
                    fte_variances.sort(key=lambda x: abs(x['variance']), reverse=True)
            
            # --- Salary by Job Class ---
            if 'Object_Key' in e_df.columns and 'JobClass_Key' in e_df.columns:
                salary_data = e_df[e_df['Object_Key'] == '51100']
                if not salary_data.empty:
                    for jc in salary_data['JobClass_Key'].unique():
                        jc_data = salary_data[salary_data['JobClass_Key'] == jc]
                        ytd_sum = jc_data[ytd_col].sum()
                        fte_sum = jc_data[fte_col].sum() if fte_col else 0
                        adj_fte_sum = jc_data[adj_fte_col].sum() if adj_fte_col else 0
                        bud_sum = jc_data[bud_col].sum() if bud_col else 0
                        enc_sum = jc_data[enc_col].sum() if enc_col else 0
                        if ytd_sum > 0:
                            jc_name = jc
                            if 'JobClass' in jc_data.columns:
                                first = jc_data['JobClass'].iloc[0]
                                if ' - ' in str(first):
                                    jc_name = str(first).split(' - ', 1)[1][:30]
                            
                            salary_by_jobclass[jc] = {
                                'name': jc_name, 'ytd': ytd_sum, 'fte': fte_sum,
                                'adj_fte': adj_fte_sum, 'budget': bud_sum,
                                'encumbrance': enc_sum,
                                'fte_var': fte_sum - adj_fte_sum,
                                'per_fte': ytd_sum / fte_sum if fte_sum > 0 else 0
                            }
    
    # --- Cash & Fund Comparison ---
    fund_data_list = []
    funds_exceed_reimbursable = []
    funds_exceed_non_reimbursable = []
    
    c_df = None
    if cash_df is not None:
        c_df = cash_df.copy()
        c_df, _ = normalize_fund_col(c_df)
        for col in c_df.columns:
            if "Line" in col:
                c_df[col] = clean_currency_series(c_df[col])
        
        col_l2 = next((c for c in c_df.columns if 'Line 2' in c), None)
        col_l5 = next((c for c in c_df.columns if 'Line 5' in c), None)
        col_l7 = next((c for c in c_df.columns if 'Line 7' in c), None)
        
        if col_l7:
            total_cash = c_df[c_df['Fund_Key'] != 'GRAND TOTAL'][col_l7].sum()
        
        # Reimbursable fund prefixes (federal flow-through, state flow-through)
        reimbursable_prefixes = ('24', '25', '26', '27', '28', '29')
        
        for _, row in c_df.iterrows():
            fund = str(row.get('Fund_Key', ''))
            if fund == 'GRAND TOTAL' or not fund:
                continue
            
            rev = row[col_l2] if col_l2 else 0
            exp = abs(row[col_l5]) if col_l5 else 0
            cash = row[col_l7] if col_l7 else 0
            net = rev - exp
            
            if cash < -0.01:
                status = 'critical'; status_text = 'Negative Cash'
            elif net < -0.01:
                status = 'warning'; status_text = 'Exp > Rev'
            else:
                status = 'healthy'; status_text = 'Healthy'
            
            fund_entry = {
                'fund': fund, 'revenue': rev, 'expenditure': exp,
                'net': net, 'cash': cash, 'status': status, 'status_text': status_text
            }
            fund_data_list.append(fund_entry)
            
            # Split exceeding funds into reimbursable vs non-reimbursable
            if net < -0.01 and exp > 0:
                is_reimbursable = fund.startswith(reimbursable_prefixes)
                entry = {
                    'fund': fund, 'revenue': rev, 'expenditure': exp, 'net': net,
                    'fund_name': exp_by_fund.get(fund, {}).get('name', fund)
                }
                if is_reimbursable:
                    funds_exceed_reimbursable.append(entry)
                else:
                    funds_exceed_non_reimbursable.append(entry)
        
        funds_exceed_non_reimbursable.sort(key=lambda x: x['net'])
        funds_exceed_reimbursable.sort(key=lambda x: x['net'])
    
    net_position = total_revenue - total_expenditure
    
    # ======================================================
    # COLLECT CONCERNS & HIGHLIGHTS FROM VALIDATIONS
    # ======================================================
    concerns_high = []
    concerns_med = []
    concerns_low = []
    highlights = []
    
    for step, findings in validation_results.items():
        for status, msg in findings:
            if status == "FLAG" and "🚩" in msg:
                # Categorize by keywords
                msg_lower = msg.lower()
                if any(k in msg_lower for k in ['material', 'exceed budget', 'negative cash', 'forbidden']):
                    concerns_high.append({'step': step, 'message': msg})
                elif any(k in msg_lower for k in ['$0 ytd', 'exceed revenue', 'negative', 'universal']):
                    concerns_med.append({'step': step, 'message': msg})
                else:
                    concerns_low.append({'step': step, 'message': msg})
            elif status == "PASS" and "✅" in msg:
                highlights.append({'step': step, 'message': msg})
    
    # ======================================================
    # BUILD ACTION ITEMS FROM CHECKLIST NOTES
    # ======================================================
    action_items = []
    if checklist_data:
        for item in checklist_data:
            if item.get('user_notes') and item['user_notes'].strip():
                notes = item['user_notes'].strip()
                step = item['step']
                # Auto-assign priority based on step area and keywords
                notes_lower = notes.lower()
                if any(k in notes_lower for k in ['material', 'audit finding', 'exceed budget', 'bar', 'negative cash', 'required', 'rfr']):
                    priority = 'High'
                elif any(k in notes_lower for k in ['confirm', 'explain', 'provide', 'clarify', 'please']):
                    priority = 'Medium'
                else:
                    priority = 'Low'
                
                # Truncate for table display
                display_note = notes[:200] + ('...' if len(notes) > 200 else '')
                action_items.append({
                    'priority': priority,
                    'step': step,
                    'action': display_note,
                    'area': item.get('review_area', '')
                })
    
    # Sort: High first, then Medium, then Low
    priority_order = {'High': 0, 'Medium': 1, 'Low': 2}
    action_items.sort(key=lambda x: priority_order.get(x['priority'], 3))
    
    # ======================================================
    # CHART DATA PREPARATION
    # ======================================================
    
    # Revenue source chart
    rev_source_labels = json.dumps([k for k, v in revenue_by_source.items() if v > 0])
    rev_source_data = json.dumps([round(v, 2) for v in revenue_by_source.values() if v > 0])
    
    # Function chart (sorted by YTD)
    sorted_functions = sorted(expenditure_by_function.items(), key=lambda x: x[1]['ytd'], reverse=True)
    func_chart_labels = json.dumps([f"{k} {v['name'][:18]}" for k, v in sorted_functions[:10]])
    func_chart_ytd = json.dumps([round(v['ytd'], 2) for k, v in sorted_functions[:10]])
    func_chart_enc = json.dumps([round(v['encumbrance'], 2) for k, v in sorted_functions[:10]])
    
    # Fund comparison chart (top 12)
    sorted_fund_compare = sorted(fund_data_list, key=lambda x: max(x['revenue'], x['expenditure']), reverse=True)[:12]
    fund_cmp_labels = json.dumps([f['fund'] for f in sorted_fund_compare])
    fund_cmp_rev = json.dumps([round(f['revenue'], 2) for f in sorted_fund_compare])
    fund_cmp_exp = json.dumps([round(f['expenditure'], 2) for f in sorted_fund_compare])
    
    # Expenditure by fund type (donut)
    fund_type_map = {
        '11': 'Operational', '12': 'Locally Authorized', '13': 'Transportation',
        '15': 'Impact Aid / Local Rev', '21': 'Food Services', '22': 'Athletics',
        '23': 'Non-Instructional', '24': 'Federal Grants', '25': 'Federal Direct',
        '27': 'State Grants', '28': 'Other State', '29': 'State/Local Grants',
        '31': 'Capital / Bonds', '41': 'GO Debt Services', '43': 'ETN Debt Services'
    }
    exp_by_type = {}
    for fund_code, fund_info in exp_by_fund.items():
        prefix = fund_code[:2]
        label = fund_type_map.get(prefix, f'Other ({prefix}xxx)')
        exp_by_type[label] = exp_by_type.get(label, 0) + fund_info['ytd']
    
    sorted_exp_type = sorted(exp_by_type.items(), key=lambda x: x[1], reverse=True)
    exp_type_labels = json.dumps([k for k, v in sorted_exp_type if v > 0])
    exp_type_data = json.dumps([round(v, 2) for k, v in sorted_exp_type if v > 0])
    
    # Salary chart
    sorted_salary = sorted(salary_by_jobclass.items(), key=lambda x: x[1]['ytd'], reverse=True)[:12]
    salary_chart_labels = json.dumps([f"{k} {v['name'][:14]}" for k, v in sorted_salary])
    salary_chart_data = json.dumps([round(v['ytd'], 2) for k, v in sorted_salary])
    
    # Program chart
    sorted_programs = sorted(expenditure_by_program.items(), key=lambda x: x[1]['ytd'], reverse=True)
    prog_with_activity = [(k, v) for k, v in sorted_programs if v['ytd'] > 0]
    prog_chart_labels = json.dumps([f"{k} {v['name'][:18]}" for k, v in prog_with_activity[:8]])
    prog_chart_data = json.dumps([round(v['ytd'], 2) for k, v in prog_with_activity[:8]])
    
    total_salary = sum(v['ytd'] for v in salary_by_jobclass.values())
    total_fte = sum(v['fte'] for v in salary_by_jobclass.values())
    total_enc_risk = sum(l['over_amount'] for l in encumbrance_risk_lines)
    
    # ======================================================
    # BUILD HTML SECTIONS
    # ======================================================
    
    # --- Concerns HTML ---
    def build_concern_card(concern_list, severity, icon, color_class):
        if not concern_list:
            return ""
        items_html = ""
        for c in concern_list[:8]:
            msg = c['message'].replace('🚩', '').replace('⚠️', '').strip()
            items_html += f'<div class="concern-card {color_class}"><div class="concern-icon">{icon}</div><div class="concern-body"><span class="concern-step">Step {c["step"]}</span>{msg}</div></div>\n'
        return items_html
    
    high_concerns_html = build_concern_card(concerns_high, 'High', '🚩', 'high')
    med_concerns_html = build_concern_card(concerns_med, 'Medium', '⚠️', 'medium')
    low_concerns_html = build_concern_card(concerns_low, 'Low', '📋', 'low')
    
    # --- Funds Exceeding Revenue ---
    def build_fund_exceed_rows(fund_list, badge_class, badge_text):
        html = ""
        for f in fund_list:
            html += f'''<tr>
                <td>{f['fund']} — {f.get('fund_name', '')}</td>
                <td class="num">${f['revenue']:,.0f}</td>
                <td class="num">${f['expenditure']:,.0f}</td>
                <td class="num negative">&minus;${abs(f['net']):,.0f}</td>
                <td><span class="badge {badge_class}">{badge_text}</span></td>
            </tr>'''
        return html
    
    non_reimb_rows = build_fund_exceed_rows(funds_exceed_non_reimbursable, 'high', 'Non-Reimbursable')
    reimb_rows = build_fund_exceed_rows(funds_exceed_reimbursable, 'low', 'Reimbursable')
    
    # --- Function Table ---
    function_rows = ""
    for func_code, fd in sorted_functions:
        pct_class = 'negative' if fd['pct_used'] > 100 else ''
        function_rows += f'''<tr>
            <td>{func_code} — {fd['name']}</td>
            <td class="num">${fd['budget']:,.0f}</td>
            <td class="num">${fd['ytd']:,.0f}</td>
            <td class="num">${fd['encumbrance']:,.0f}</td>
            <td class="num">${fd['available']:,.0f}</td>
            <td class="num {pct_class}">{fd['pct_used']:.1f}%</td>
        </tr>'''
    
    # --- Encumbrance Risk Table ---
    enc_risk_rows = ""
    for line in encumbrance_risk_lines:
        enc_risk_rows += f'''<tr>
            <td>{line['fund']}</td>
            <td>{line['function']}</td>
            <td>{line['object']}</td>
            <td class="num">${line['budget']:,.0f}</td>
            <td class="num">${line['ytd']:,.0f}</td>
            <td class="num">${line['encumbered']:,.0f}</td>
            <td class="num">${line['total_committed']:,.0f}</td>
            <td class="num negative">${line['over_amount']:,.0f}</td>
        </tr>'''
    
    # --- Burn Rate Tables ---
    burn_over_rows = ""
    for line in burn_over_pace:
        burn_over_rows += f'''<tr>
            <td>{line['fund']}</td>
            <td>{line['function']}</td>
            <td>{line['object']}</td>
            <td class="num">${line['budget']:,.0f}</td>
            <td class="num">${line['ytd']:,.0f}</td>
            <td class="num negative">{line['burn_pct']:.1f}%</td>
        </tr>'''
    
    burn_under_rows = ""
    for line in burn_under_pace:
        burn_under_rows += f'''<tr>
            <td>{line['fund']}</td>
            <td>{line['function']}</td>
            <td>{line['object']}</td>
            <td class="num">${line['budget']:,.0f}</td>
            <td class="num">$0</td>
            <td class="num">0.0%</td>
        </tr>'''
    
    # --- FTE Variance Table ---
    fte_var_rows = ""
    for fv in fte_variances:
        var_class = 'negative' if fv['variance'] < 0 else 'positive'
        var_sign = '' if fv['variance'] < 0 else '+'
        status_badge = 'high' if fv['status'] == 'overstaffed' else 'medium'
        fte_var_rows += f'''<tr>
            <td>{fv['jobclass_code']} — {fv['jobclass']}</td>
            <td class="num">{fv['actual_fte']:.1f}</td>
            <td class="num">{fv['budgeted_fte']:.1f}</td>
            <td class="num {var_class}">{var_sign}{fv['variance']:.1f}</td>
            <td><span class="badge {status_badge}">{fv['status'].title()}</span></td>
            <td>{fv['note']}</td>
        </tr>'''
    
    # --- Salary Table ---
    salary_rows = ""
    for jc, jd in sorted_salary:
        fte_display = f"{jd['fte']:.1f}" if jd['fte'] > 0 else "—"
        adj_display = f"{jd['adj_fte']:.1f}" if jd['adj_fte'] > 0 else "—"
        var = jd['fte_var']
        var_class = 'negative' if var < -1 else ('positive' if var > 1 else '')
        var_sign = '' if var < 0 else '+'
        var_display = f"{var_sign}{var:.1f}" if jd['adj_fte'] > 0 else "—"
        avg_sal = jd['per_fte']  # already calculated as ytd/fte
        avg_sal_display = f"${avg_sal:,.0f}" if avg_sal > 0 else "—"
        salary_rows += f'''<tr>
            <td>{jc} — {jd['name']}</td>
            <td class="num">${jd['budget']:,.0f}</td>
            <td class="num">${jd['ytd']:,.0f}</td>
            <td class="num">${jd['encumbrance']:,.0f}</td>
            <td class="num">{fte_display}</td>
            <td class="num">{adj_display}</td>
            <td class="num {var_class}">{var_display}</td>
            <td class="num">{avg_sal_display}</td>
        </tr>'''
    
    # --- Program Table ---
    program_rows = ""
    for prog_code, pd_item in sorted_programs:
        if pd_item['ytd'] > 0 or pd_item['budget'] > 0:
            pct_class = 'negative' if pd_item['pct_used'] > 100 else ''
            program_rows += f'''<tr>
                <td>{prog_code} — {pd_item['name']}</td>
                <td class="num">${pd_item['ytd']:,.0f}</td>
                <td class="num">${pd_item['encumbrance']:,.0f}</td>
                <td class="num">${pd_item['budget']:,.0f}</td>
                <td class="num {pct_class}">{pd_item['pct_used']:.1f}%</td>
            </tr>'''
    
    # --- Compliance Highlights ---
    highlights_html = ""
    for h in highlights[:15]:
        msg = h['message'].replace('✅', '').strip()
        highlights_html += f'<div class="compliance-item"><div class="check">✓</div><div>{msg}</div></div>\n'
    
    # --- Action Items Table ---
    action_rows = ""
    for ai in action_items:
        badge_class = {'High': 'high', 'Medium': 'medium', 'Low': 'low'}.get(ai['priority'], 'low')
        due = {'High': high_due, 'Medium': med_due, 'Low': low_due}.get(ai['priority'], low_due)
        action_rows += f'''<tr>
            <td><span class="badge {badge_class}">{ai['priority']}</span></td>
            <td>{ai['action']}</td>
            <td>Step {ai['step']}</td>
            <td>{due}</td>
        </tr>'''
    
    # --- Review Notes / Audit Findings ---
    audit_notes_html = ""
    findings_html = ""
    if checklist_data:
        for item in checklist_data:
            if item.get('user_notes') and item['user_notes'].strip():
                step = item['step']
                notes = item['user_notes'].strip().replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br>')
                check = item['check'][:60]
                area = item['review_area']
                
                if step == 62:
                    audit_notes_html = notes
                else:
                    findings_html += f'''<div class="finding-item">
                        <div class="finding-header">
                            <span class="step-badge">Step {step}</span>
                            <span class="finding-area">{area}</span>
                            <span class="finding-check">{check}</span>
                        </div>
                        <div class="finding-notes">{notes}</div>
                    </div>\n'''
    
    # ======================================================
    # FULL HTML OUTPUT
    # ======================================================
    
    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{entity_name} — {fiscal_year} {quarter_label} Actuals Analysis</title>
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;600;700;800&family=Source+Sans+3:wght@300;400;500;600;700&display=swap" rel="stylesheet">
<script src="https://cdnjs.cloudflare.com/ajax/libs/Chart.js/4.4.1/chart.umd.min.js"></script>
<style>
  :root {{
    --bg: #faf9f7;
    --card: #ffffff;
    --text: #1a1a1a;
    --muted: #6b7280;
    --accent: #1e3a5f;
    --accent-light: #2d5a8e;
    --border: #e5e2dc;
    --red: #b91c1c;
    --red-bg: #fef2f2;
    --red-border: #fecaca;
    --orange: #c2410c;
    --orange-bg: #fff7ed;
    --orange-border: #fed7aa;
    --green: #166534;
    --green-bg: #f0fdf4;
    --green-border: #bbf7d0;
    --blue-bg: #eff6ff;
    --blue-border: #bfdbfe;
    --yellow: #a16207;
    --yellow-bg: #fefce8;
    --yellow-border: #fef08a;
    --shadow: 0 1px 3px rgba(0,0,0,0.06), 0 1px 2px rgba(0,0,0,0.04);
    --shadow-md: 0 4px 6px -1px rgba(0,0,0,0.07), 0 2px 4px -2px rgba(0,0,0,0.05);
  }}
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{
    font-family: 'Source Sans 3', 'Source Sans Pro', system-ui, sans-serif;
    background: var(--bg);
    color: var(--text);
    line-height: 1.6;
    -webkit-font-smoothing: antialiased;
  }}
  .container {{ max-width: 1200px; margin: 0 auto; padding: 0 24px; }}
  header {{
    background: linear-gradient(135deg, #1e3a5f 0%, #2d5a8e 50%, #1e3a5f 100%);
    color: white;
    padding: 48px 0 40px;
    border-bottom: 4px solid #c9a96e;
  }}
  header .entity-name {{
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 2rem; font-weight: 700; letter-spacing: -0.02em; margin-bottom: 4px;
  }}
  header .report-title {{
    font-size: 1.1rem; font-weight: 300; opacity: 0.9; letter-spacing: 0.03em;
    text-transform: uppercase; margin-bottom: 16px;
  }}
  header .meta {{ font-size: 0.88rem; opacity: 0.75; display: flex; gap: 24px; flex-wrap: wrap; }}
  .section {{ margin: 40px 0; }}
  .section-title {{
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 1.45rem; font-weight: 700; color: var(--accent);
    margin-bottom: 6px; padding-bottom: 8px; border-bottom: 2px solid var(--border);
  }}
  .section-subtitle {{ color: var(--muted); font-size: 0.9rem; margin-bottom: 20px; }}
  .summary-grid {{
    display: grid; grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
    gap: 20px; margin-top: 20px;
  }}
  .summary-card {{
    background: var(--card); border-radius: 8px; box-shadow: var(--shadow-md);
    border: 1px solid var(--border); padding: 28px 24px; text-align: center;
    position: relative; overflow: hidden;
  }}
  .summary-card::before {{ content: ''; position: absolute; top: 0; left: 0; right: 0; height: 4px; }}
  .summary-card.revenue::before {{ background: #166534; }}
  .summary-card.expenditure::before {{ background: #b91c1c; }}
  .summary-card.cash::before {{ background: #1e3a5f; }}
  .summary-card.net::before {{ background: #c2410c; }}
  .summary-card.enc::before {{ background: #7c3aed; }}
  .summary-card .label {{
    font-size: 0.82rem; text-transform: uppercase; letter-spacing: 0.08em;
    color: var(--muted); font-weight: 600; margin-bottom: 6px;
  }}
  .summary-card .value {{
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 1.65rem; font-weight: 700; color: var(--text);
  }}
  .summary-card .sub {{ font-size: 0.82rem; color: var(--muted); margin-top: 4px; }}
  
  .concern-card {{
    border-radius: 8px; padding: 14px 18px; border-left: 4px solid;
    display: flex; gap: 12px; align-items: flex-start; margin-bottom: 10px;
  }}
  .concern-card.high {{ background: var(--red-bg); border-color: var(--red); }}
  .concern-card.medium {{ background: var(--orange-bg); border-color: var(--orange); }}
  .concern-card.low {{ background: var(--yellow-bg); border-color: var(--yellow); }}
  .concern-icon {{ font-size: 1.1rem; flex-shrink: 0; margin-top: 2px; }}
  .concern-body {{ font-size: 0.88rem; color: #374151; line-height: 1.5; }}
  .concern-step {{
    display: inline-block; background: rgba(0,0,0,0.08); border-radius: 4px;
    padding: 1px 6px; font-size: 0.75rem; font-weight: 700; margin-right: 6px;
  }}
  
  .table-wrap {{ overflow-x: auto; border-radius: 8px; border: 1px solid var(--border); box-shadow: var(--shadow); }}
  table {{ width: 100%; border-collapse: collapse; font-size: 0.88rem; }}
  thead th {{
    background: #f8f7f5; font-weight: 700; text-align: left; padding: 12px 14px;
    border-bottom: 2px solid var(--border); white-space: nowrap; font-size: 0.82rem;
    text-transform: uppercase; letter-spacing: 0.04em; color: var(--accent);
  }}
  tbody td {{ padding: 10px 14px; border-bottom: 1px solid #f0eeea; vertical-align: top; }}
  tbody tr:hover {{ background: #f9f8f6; }}
  tbody tr:last-child td {{ border-bottom: none; }}
  td.num {{ text-align: right; font-variant-numeric: tabular-nums; }}
  td.negative, .negative {{ color: var(--red); font-weight: 600; }}
  td.positive, .positive {{ color: var(--green); font-weight: 600; }}
  
  .badge {{
    display: inline-block; padding: 2px 10px; border-radius: 12px;
    font-size: 0.75rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.03em;
  }}
  .badge.high {{ background: #fecaca; color: var(--red); }}
  .badge.medium {{ background: #fed7aa; color: var(--orange); }}
  .badge.low {{ background: #fef08a; color: var(--yellow); }}
  .badge.ok {{ background: #bbf7d0; color: var(--green); }}
  
  .compliance-item {{
    display: flex; gap: 10px; align-items: flex-start;
    padding: 12px 16px; background: var(--green-bg); border-radius: 8px;
    border: 1px solid var(--green-border); margin-bottom: 8px; font-size: 0.88rem;
  }}
  .compliance-item .check {{ color: var(--green); font-size: 1.1rem; flex-shrink: 0; }}
  
  .chart-container {{
    position: relative; padding: 16px; background: var(--card);
    border-radius: 8px; border: 1px solid var(--border); box-shadow: var(--shadow);
  }}
  .chart-row {{ display: grid; grid-template-columns: 1fr 1fr; gap: 24px; align-items: start; }}
  
  .finding-item {{
    border: 1px solid var(--border); border-radius: 8px;
    margin-bottom: 10px; overflow: hidden;
  }}
  .finding-header {{
    display: flex; align-items: center; gap: 10px;
    padding: 10px 14px; background: #f8f7f5;
  }}
  .step-badge {{
    background: var(--accent); color: white; padding: 2px 8px;
    border-radius: 4px; font-size: 0.72rem; font-weight: 700; white-space: nowrap;
  }}
  .finding-area {{ font-size: 0.78rem; color: var(--muted); }}
  .finding-check {{ flex: 1; font-size: 0.88rem; font-weight: 500; }}
  .finding-notes {{ padding: 12px 14px; font-size: 0.85rem; color: #374151; line-height: 1.55; }}
  
  .audit-content {{
    background: #f8f7f5; padding: 16px 20px; border-radius: 8px;
    font-size: 0.88rem; line-height: 1.6;
  }}
  
  .subsection-title {{
    font-size: 1rem; font-weight: 700; color: var(--accent);
    margin: 20px 0 10px; padding-bottom: 4px; border-bottom: 1px solid var(--border);
  }}
  .info-box {{
    padding: 14px 18px; border-radius: 8px; font-size: 0.88rem;
    margin-bottom: 16px; border: 1px solid;
  }}
  .info-box.blue {{ background: var(--blue-bg); border-color: var(--blue-border); color: #1e3a5f; }}
  .info-box.red {{ background: var(--red-bg); border-color: var(--red-border); color: var(--red); }}
  .info-box.yellow {{ background: var(--yellow-bg); border-color: var(--yellow-border); color: var(--yellow); }}

  footer {{
    background: var(--accent); color: white; text-align: center;
    padding: 24px; margin-top: 60px; font-size: 0.82rem;
  }}
  footer .fine {{ margin-top: 8px; opacity: 0.7; font-size: 0.75rem; font-style: italic; }}

  @media (max-width: 768px) {{
    .chart-row {{ grid-template-columns: 1fr; }}
    header .entity-name {{ font-size: 1.5rem; }}
    .summary-card .value {{ font-size: 1.3rem; }}
    .summary-grid {{ grid-template-columns: repeat(2, 1fr); }}
  }}
  @media print {{
    body {{ background: white; }}
    .chart-container, .section {{ break-inside: avoid; }}
    header {{ background: var(--accent) !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
  }}
</style>
</head>
<body>

<header>
  <div class="container">
    <div class="entity-name">{entity_name}</div>
    <div class="report-title">{fiscal_year} {quarter_label} Actuals Analysis Dashboard</div>
    <div class="meta">
      <span>Review Date: {review_date}</span>
      <span>NM Public Education Department / School Budget Bureau</span>
    </div>
  </div>
</header>

<div class="container">

<!-- 1. EXECUTIVE SUMMARY -->
<div class="section">
  <h2 class="section-title">Executive Summary</h2>
  <div class="summary-grid">
    <div class="summary-card revenue">
      <div class="label">Revenue YTD</div>
      <div class="value">${total_revenue:,.0f}</div>
      <div class="sub">Excludes cash balance lines</div>
    </div>
    <div class="summary-card expenditure">
      <div class="label">Expenditure YTD</div>
      <div class="value">${total_expenditure:,.0f}</div>
      <div class="sub">${total_encumbrance:,.0f} encumbered</div>
    </div>
    <div class="summary-card cash">
      <div class="label">Cash Balance</div>
      <div class="value">${total_cash:,.0f}</div>
      <div class="sub">Line 7 — all funds</div>
    </div>
    <div class="summary-card net">
      <div class="label">Net Rev &minus; Exp</div>
      <div class="value" style="color:{'var(--green)' if net_position >= 0 else 'var(--red)'};">{'&minus;' if net_position < 0 else ''}${abs(net_position):,.0f}</div>
      <div class="sub">{'Revenue exceeds expenditures' if net_position >= 0 else 'Expenditures exceed revenue'}</div>
    </div>
    <div class="summary-card enc">
      <div class="label">Total Budget</div>
      <div class="value">${total_budget:,.0f}</div>
      <div class="sub">{((total_expenditure + total_encumbrance) / total_budget * 100) if total_budget > 0 else 0:.1f}% committed</div>
    </div>
  </div>
</div>

<!-- 2. KEY CONCERNS -->
{'<div class="section"><h2 class="section-title">Key Concerns Requiring Attention</h2><p class="section-subtitle">Issues grouped by severity — High (red), Medium (orange), Low (yellow)</p>' + high_concerns_html + med_concerns_html + low_concerns_html + '</div>' if (high_concerns_html or med_concerns_html or low_concerns_html) else ''}

<!-- 3. REVENUE VS EXPENDITURE -->
<div class="section">
  <h2 class="section-title">Revenue vs. Expenditure Analysis</h2>
  <p class="section-subtitle">Top 12 funds by total financial activity (YTD)</p>
  <div class="chart-container" style="height:440px;">
    <canvas id="fundCompareChart"></canvas>
  </div>
</div>

<!-- 4. FUNDS WHERE EXP > REV (Split) -->
{'<div class="section"><h2 class="section-title">Funds Where Expenditures Exceed Revenue</h2>' + (f'<div class="subsection-title">Non-Reimbursable Funds (Requires Explanation)</div><div class="info-box red"><strong>{len(funds_exceed_non_reimbursable)} fund(s)</strong> have expenditures exceeding revenue in non-reimbursable categories. These require immediate attention.</div><div class="table-wrap"><table><thead><tr><th>Fund</th><th style="text-align:right">Revenue</th><th style="text-align:right">Expenditure</th><th style="text-align:right">Net</th><th>Type</th></tr></thead><tbody>{non_reimb_rows}</tbody></table></div>' if non_reimb_rows else '') + (f'<div class="subsection-title" style="margin-top:24px;">Reimbursable Funds (Expected Behavior)</div><p class="section-subtitle">Federal/state flow-through funds may carry negative net until reimbursed.</p><div class="table-wrap"><table><thead><tr><th>Fund</th><th style="text-align:right">Revenue</th><th style="text-align:right">Expenditure</th><th style="text-align:right">Net</th><th>Type</th></tr></thead><tbody>{reimb_rows}</tbody></table></div>' if reimb_rows else '') + '</div>' if (non_reimb_rows or reimb_rows) else ''}

<!-- 5. EXPENDITURE BY FUND TYPE (Donut + Table) -->
<div class="section">
  <h2 class="section-title">Expenditure Distribution by Fund Type</h2>
  <p class="section-subtitle">Total YTD expenditures: ${total_expenditure:,.0f}</p>
  <div class="chart-row">
    <div class="chart-container" style="height:380px;">
      <canvas id="expTypeDonut"></canvas>
    </div>
    <div class="table-wrap">
      <table>
        <thead><tr><th>Fund Type</th><th style="text-align:right">YTD Spend</th><th style="text-align:right">% of Total</th></tr></thead>
        <tbody>
          {''.join(f'<tr><td>{k}</td><td class="num">${v:,.0f}</td><td class="num">{(v/total_expenditure*100) if total_expenditure > 0 else 0:.1f}%</td></tr>' for k, v in sorted_exp_type if v > 0)}
        </tbody>
      </table>
    </div>
  </div>
</div>

<!-- 6. EXPENDITURE BY FUNCTION -->
<div class="section">
  <h2 class="section-title">Expenditure by Function</h2>
  <p class="section-subtitle">Functional breakdown — budget utilization through {quarter_label}</p>
  <div class="chart-container" style="height:420px;margin-bottom:20px;">
    <canvas id="functionChart"></canvas>
  </div>
  <div class="table-wrap">
    <table>
      <thead><tr><th>Function</th><th style="text-align:right">Budget</th><th style="text-align:right">YTD Actuals</th><th style="text-align:right">Encumbered</th><th style="text-align:right">Available</th><th style="text-align:right">% Used (Act+Enc)</th></tr></thead>
      <tbody>{function_rows}</tbody>
    </table>
  </div>
</div>

<!-- 7. ENCUMBRANCE RISK -->
{'<div class="section"><h2 class="section-title">Encumbrance Risk Analysis</h2><p class="section-subtitle">Lines where Actuals + Encumbrances exceed Adjusted Budget — total over-committed: <strong style="color:var(--red);">${total_enc_risk:,.0f}</strong></p><div class="info-box red">These represent committed over-expenditures. If encumbrances convert to actuals, these lines will be over budget. BARs should be submitted to restore budget authority.</div><div class="table-wrap"><table><thead><tr><th>Fund</th><th>Function</th><th>Object</th><th style="text-align:right">Budget</th><th style="text-align:right">Actuals</th><th style="text-align:right">Encumbered</th><th style="text-align:right">Total Committed</th><th style="text-align:right">Over Amount</th></tr></thead><tbody>' + enc_risk_rows + '</tbody></table></div></div>' if enc_risk_rows else ''}

<!-- 8. BURN RATE -->
{'<div class="section"><h2 class="section-title">Burn Rate / Pace Analysis</h2><p class="section-subtitle">At ' + quarter_label + ', expected actuals burn rate is ~25%. Flagging outliers.</p>' + (f'<div class="subsection-title">Over-Pace (Burn Rate &gt; 40% at {quarter_label})</div><p class="section-subtitle">These lines are spending faster than expected and may overspend by year-end.</p><div class="table-wrap"><table><thead><tr><th>Fund</th><th>Function</th><th>Object</th><th style="text-align:right">Budget</th><th style="text-align:right">YTD</th><th style="text-align:right">Burn %</th></tr></thead><tbody>{burn_over_rows}</tbody></table></div>' if burn_over_rows else '') + (f'<div class="subsection-title" style="margin-top:24px;">Under-Pace / No Activity (Budget &gt; $50k, $0 YTD)</div><p class="section-subtitle">Significant budget authority with no actuals — may indicate delayed programs or coding issues.</p><div class="table-wrap"><table><thead><tr><th>Fund</th><th>Function</th><th>Object</th><th style="text-align:right">Budget</th><th style="text-align:right">YTD</th><th style="text-align:right">Burn %</th></tr></thead><tbody>{burn_under_rows}</tbody></table></div>' if burn_under_rows else '') + '</div>' if (burn_over_rows or burn_under_rows) else ''}

<!-- 9. FTE VARIANCE -->
{'<div class="section"><h2 class="section-title">FTE Variance Analysis</h2><p class="section-subtitle">Object 51100 staffing — flagging variances of 2.0+ FTE between actuals and budget</p><div class="table-wrap"><table><thead><tr><th>Job Class</th><th style="text-align:right">Actual FTE</th><th style="text-align:right">Budgeted FTE</th><th style="text-align:right">Variance</th><th>Status</th><th>Note</th></tr></thead><tbody>' + fte_var_rows + '</tbody></table></div></div>' if fte_var_rows else ''}

<!-- 10. SALARY BY JOB CLASS -->
{'<div class="section"><h2 class="section-title">Salary by Job Class (Object 51100)</h2><p class="section-subtitle">Total salary: <strong>${total_salary:,.0f}</strong> | Total FTE: <strong>{total_fte:.1f}</strong></p><div class="chart-container" style="height:380px;margin-bottom:20px;"><canvas id="salaryChart"></canvas></div><div class="table-wrap"><table><thead><tr><th>Job Class</th><th style="text-align:right">Budget</th><th style="text-align:right">YTD Salary</th><th style="text-align:right">Encumbered</th><th style="text-align:right">Actual FTE</th><th style="text-align:right">Budgeted FTE</th><th style="text-align:right">FTE &Delta;</th><th style="text-align:right">Avg Salary</th></tr></thead><tbody>' + salary_rows + '</tbody></table></div></div>' if salary_rows else ''}

<!-- 11. PROGRAM-LEVEL SPENDING -->
{'<div class="section"><h2 class="section-title">Program-Level Spending</h2><p class="section-subtitle">Expenditure breakdown by program code</p><div class="chart-row"><div class="chart-container" style="height:320px;"><canvas id="programChart"></canvas></div><div class="table-wrap"><table><thead><tr><th>Program</th><th style="text-align:right">YTD Actuals</th><th style="text-align:right">Encumbered</th><th style="text-align:right">Budget</th><th style="text-align:right">% Used</th></tr></thead><tbody>' + program_rows + '</tbody></table></div></div></div>' if program_rows else ''}

<!-- 12. COMPLIANCE HIGHLIGHTS -->
{'<div class="section"><h2 class="section-title">Compliance Highlights</h2><p class="section-subtitle">Items that passed automated and manual checks</p>' + highlights_html + '</div>' if highlights_html else ''}

<!-- 13. PRIOR YEAR AUDIT -->
{'<div class="section"><h2 class="section-title">Prior Year Audit Findings</h2><div class="audit-content">' + audit_notes_html + '</div></div>' if audit_notes_html else ''}

<!-- 14. ACTION ITEMS -->
{'<div class="section"><h2 class="section-title">Action Items for District Response</h2><p class="section-subtitle">Prioritized items — response due dates based on priority level</p><div class="table-wrap"><table><thead><tr><th style="width:70px;">Priority</th><th>Action Item</th><th style="width:80px;">Ref</th><th style="width:140px;">Response Due</th></tr></thead><tbody>' + action_rows + '</tbody></table></div></div>' if action_rows else ''}

<!-- 15. REVIEW NOTES -->
{'<div class="section"><h2 class="section-title">Detailed Review Notes</h2>' + findings_html + '</div>' if findings_html else ''}

</div><!-- end container -->

<!-- 16. DATA SOURCES & METHODOLOGY -->
<footer>
  <div>NM Public Education Department — School Budget Bureau — {fiscal_year} {quarter_label} Quarterly Actuals Review</div>
  <div style="margin-top:6px;">Report generated {review_date} | Confidential — For Official Use Only</div>
  <div class="fine">
    Data Sources: OBMS Revenue Actuals Report, OBMS Expenditure Actuals Report, Cash Report (Excel/Summary Tab), Bank Statement(s).
    This review is conducted pursuant to NMAC 6.20.2 and does not constitute an audit.
  </div>
</footer>

<script>
Chart.defaults.font.family = "'Source Sans 3', system-ui, sans-serif";
Chart.defaults.font.size = 12;
Chart.defaults.color = '#374151';

const palette = ['#1e3a5f','#b91c1c','#c9a96e','#3b7dd8','#0d9488','#ea580c','#166534','#7c3aed','#db2777','#4338ca','#059669','#d97706','#0891b2','#e11d48','#64748b'];

// Fund Comparison Bar
new Chart(document.getElementById('fundCompareChart').getContext('2d'), {{
  type: 'bar',
  data: {{
    labels: {fund_cmp_labels},
    datasets: [
      {{ label: 'Revenue YTD', data: {fund_cmp_rev}, backgroundColor: '#3b7dd8', borderRadius: 3 }},
      {{ label: 'Expenditure YTD', data: {fund_cmp_exp}, backgroundColor: '#b91c1c', borderRadius: 3 }}
    ]
  }},
  options: {{
    responsive: true, maintainAspectRatio: false,
    plugins: {{
      legend: {{ position: 'top', labels: {{ padding: 16, usePointStyle: true, pointStyle: 'rectRounded' }} }},
      tooltip: {{ callbacks: {{ label: ctx => ctx.dataset.label + ': $' + ctx.parsed.y.toLocaleString() }} }}
    }},
    scales: {{
      x: {{ grid: {{ display: false }}, ticks: {{ maxRotation: 45, font: {{ size: 11 }} }} }},
      y: {{ grid: {{ color: '#f0eeea' }}, ticks: {{ callback: v => v >= 1e6 ? '$' + (v/1e6).toFixed(1) + 'M' : '$' + (v/1e3).toFixed(0) + 'k' }} }}
    }}
  }}
}});

// Expenditure Type Donut
new Chart(document.getElementById('expTypeDonut').getContext('2d'), {{
  type: 'doughnut',
  data: {{
    labels: {exp_type_labels},
    datasets: [{{ data: {exp_type_data}, backgroundColor: palette.slice(0, {len([v for v in sorted_exp_type if v[1] > 0])}), borderWidth: 2, borderColor: '#faf9f7' }}]
  }},
  options: {{
    responsive: true, maintainAspectRatio: false, cutout: '55%',
    plugins: {{
      legend: {{ position: 'bottom', labels: {{ padding: 10, usePointStyle: true, pointStyle: 'circle', font: {{ size: 11 }} }} }},
      tooltip: {{ callbacks: {{ label: ctx => {{ const pct = ((ctx.parsed / {total_expenditure if total_expenditure > 0 else 1}) * 100).toFixed(1); return ctx.label + ': $' + ctx.parsed.toLocaleString() + ' (' + pct + '%)'; }} }} }}
    }}
  }}
}});

// Function Stacked Bar
new Chart(document.getElementById('functionChart').getContext('2d'), {{
  type: 'bar',
  data: {{
    labels: {func_chart_labels},
    datasets: [
      {{ label: 'YTD Actuals', data: {func_chart_ytd}, backgroundColor: '#1e3a5f', borderRadius: 3 }},
      {{ label: 'Encumbered', data: {func_chart_enc}, backgroundColor: '#c9a96e', borderRadius: 3 }}
    ]
  }},
  options: {{
    indexAxis: 'y', responsive: true, maintainAspectRatio: false,
    plugins: {{ legend: {{ position: 'top', labels: {{ padding: 16, usePointStyle: true, pointStyle: 'rectRounded' }} }},
      tooltip: {{ callbacks: {{ label: ctx => ctx.dataset.label + ': $' + ctx.parsed.x.toLocaleString() }} }} }},
    scales: {{
      x: {{ stacked: true, grid: {{ color: '#f0eeea' }}, ticks: {{ callback: v => v >= 1e6 ? '$' + (v/1e6).toFixed(0) + 'M' : '$' + (v/1e3).toFixed(0) + 'k' }} }},
      y: {{ stacked: true, grid: {{ display: false }}, ticks: {{ font: {{ size: 11 }} }} }}
    }}
  }}
}});

// Salary Chart
{f"""new Chart(document.getElementById('salaryChart').getContext('2d'), {{
  type: 'bar',
  data: {{
    labels: {salary_chart_labels},
    datasets: [{{ data: {salary_chart_data}, backgroundColor: '#1e3a5f', borderRadius: 3 }}]
  }},
  options: {{
    indexAxis: 'y', responsive: true, maintainAspectRatio: false,
    plugins: {{ legend: {{ display: false }},
      tooltip: {{ callbacks: {{ label: ctx => '$' + ctx.parsed.x.toLocaleString() }} }} }},
    scales: {{
      x: {{ grid: {{ color: '#f0eeea' }}, ticks: {{ callback: v => '$' + (v/1e3).toFixed(0) + 'k' }} }},
      y: {{ grid: {{ display: false }}, ticks: {{ font: {{ size: 11 }} }} }}
    }}
  }}
}});""" if salary_by_jobclass else ""}

// Program Chart
{f"""new Chart(document.getElementById('programChart').getContext('2d'), {{
  type: 'doughnut',
  data: {{
    labels: {prog_chart_labels},
    datasets: [{{ data: {prog_chart_data}, backgroundColor: palette.slice(0, {len(prog_with_activity[:8])}), borderWidth: 2, borderColor: '#faf9f7' }}]
  }},
  options: {{
    responsive: true, maintainAspectRatio: false, cutout: '50%',
    plugins: {{
      legend: {{ position: 'bottom', labels: {{ padding: 10, usePointStyle: true, pointStyle: 'circle', font: {{ size: 11 }} }} }},
      tooltip: {{ callbacks: {{ label: ctx => ctx.label + ': $' + ctx.parsed.toLocaleString() }} }}
    }}
  }}
}});""" if prog_with_activity else ""}
</script>
</body>
</html>'''
    
    return html

# ---------- EXPORT FUNCTIONS ----------

def add_findings_table_to_doc(doc, table_data: Dict):
    """Add a formatted table to a Word document from table_findings data."""
    title = table_data.get('title', 'Details')
    df = table_data['data']
    
    if df.empty:
        return
    
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)
    
    num_cols = len(df.columns)
    num_rows = len(df) + 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    header_row = table.rows[0]
    for i, col_name in enumerate(df.columns):
        cell = header_row.cells[i]
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A2744"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    for row_idx, (_, row) in enumerate(df.iterrows()):
        table_row = table.rows[row_idx + 1]
        for col_idx, col_name in enumerate(df.columns):
            cell = table_row.cells[col_idx]
            cell_val = str(row[col_name])
            cell.text = cell_val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if cell_val.startswith('-') or cell_val.startswith('($') or cell_val.startswith('-$'):
                        run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
        
        if row_idx % 2 == 1:
            for cell in table_row.cells:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7FAFC"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()


def export_findings_memo(checklist_data: List[Dict], entity_name: str, analysis_summary: Dict = None, table_findings: Dict = None) -> BytesIO:
    doc = Document()
    doc.add_heading('Actuals Analysis Findings Memo', 0)
    doc.add_paragraph(f'Entity: {entity_name}')
    doc.add_paragraph(f'Review Date: {datetime.now().strftime("%B %d, %Y")}')
    
    # Add Analysis Summary if available
    if analysis_summary:
        doc.add_heading('Executive Summary', level=1)
        
        # Statistics
        if analysis_summary.get('statistics'):
            stats = analysis_summary['statistics']
            p = doc.add_paragraph()
            if 'total_revenue' in stats:
                p.add_run(f"Total Revenue YTD: ${stats['total_revenue']:,.2f}\n")
            if 'total_expenditure' in stats:
                p.add_run(f"Total Expenditure YTD: ${stats['total_expenditure']:,.2f}\n")
            if 'total_cash' in stats:
                p.add_run(f"Total Cash Balance: ${stats['total_cash']:,.2f}\n")
        
        # Key Concerns
        if analysis_summary.get('concerns'):
            doc.add_heading('Key Concerns', level=2)
            for concern in analysis_summary['concerns'][:10]:  # Top 10 concerns
                doc.add_paragraph(concern, style='List Bullet')
        
        # Compliance Highlights
        if analysis_summary.get('highlights'):
            doc.add_heading('Compliance Highlights', level=2)
            for highlight in analysis_summary['highlights'][:10]:  # Top 10 highlights
                doc.add_paragraph(highlight, style='List Bullet')
    
    # --- Table Findings ---
    if table_findings:
        doc.add_heading('Automated Findings Detail Tables', level=1)
        for step in sorted(table_findings.keys()):
            tf = table_findings[step]
            doc.add_heading(f'Step {step}', level=2)
            add_findings_table_to_doc(doc, tf)
    
    doc.add_paragraph()

    # Original checklist findings
    doc.add_heading('Detailed Checklist Findings', level=1)
    
    review_areas = {}
    for item in checklist_data:
        if item.get('user_notes') or not item.get('completed'):
            area = item['review_area']
            if area not in review_areas: review_areas[area] = []
            review_areas[area].append(item)

    if not review_areas:
        doc.add_paragraph("No findings or incomplete items to report.")
    else:
        for area_name, items in review_areas.items():
            doc.add_heading(area_name, level=2)
            for item in items:
                doc.add_heading(f"Step {item['step']}: {item['check']}", level=3)
                status = "✓ Completed" if item['completed'] else "⚠ Incomplete"
                p = doc.add_paragraph()
                p.add_run(f"Status: {status}\n").bold = True
                if item['user_notes']:
                    p.add_run("\nFindings:\n").bold = True
                    doc.add_paragraph(item['user_notes'], style='List Bullet')
                doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_checklist_tracker(checklist_data: List[Dict]) -> BytesIO:
    df = pd.DataFrame([{
            "Step": item['step'],
            "Review Area": item['review_area'],
            "Check": item['check'],
            "Status": "✓" if item['completed'] else "✗",
            "Notes": item['user_notes']
        } for item in checklist_data])

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Checklist', index=False)
    buffer.seek(0)
    return buffer

# ---------- SESSION STATE ----------

if 'checklist_data' not in st.session_state:
    st.session_state.checklist_data = load_official_checklist()
if 'cash_df' not in st.session_state: st.session_state.cash_df = None
if 'revenue_df' not in st.session_state: st.session_state.revenue_df = None
if 'expenditure_df' not in st.session_state: st.session_state.expenditure_df = None
if 'entity_name' not in st.session_state: st.session_state.entity_name = ""
if 'validation_results' not in st.session_state: st.session_state.validation_results = {}
if 'table_findings' not in st.session_state: st.session_state.table_findings = {}

if 'notes_by_step' not in st.session_state:
    st.session_state.notes_by_step = {}
    for item in st.session_state.checklist_data:
        st.session_state.notes_by_step[item['step']] = item.get('user_notes', "")

if 'welcome_dismissed' not in st.session_state:
    st.session_state.welcome_dismissed = False

# ---------- NOTES PERSISTENCE CALLBACK ----------
def save_note_callback(step_id: int):
    '''Callback to immediately save notes when changed - fixes notes not saving issue'''
    key = f"n_{step_id}"
    if key in st.session_state:
        st.session_state.notes_by_step[step_id] = st.session_state[key]
        # Also update checklist_data
        for item in st.session_state.checklist_data:
            if item['step'] == step_id:
                item['user_notes'] = st.session_state[key]
                break
# ---------- RENDER TABLE FINDINGS IN STREAMLIT ----------
def render_findings_table(step_id: int, table_findings: Dict):
    """Render a table finding as a styled dataframe in Streamlit."""
    if step_id in table_findings:
        tf = table_findings[step_id]
        title = tf.get('title', '')
        df = tf['data']
        if not df.empty:
            st.caption(f"📋 {title}")
            st.dataframe(df, use_container_width=True, hide_index=True)

@st.dialog("Welcome to Actuals Analysis & Compliance", width="large")
def render_welcome_modal():
    """
    One-time welcome walkthrough for new users.
    Uses Streamlit's native dialog decorator (st.dialog).
    """

    # Step tracker inside the modal
    if 'welcome_step' not in st.session_state:
        st.session_state.welcome_step = 1

    step = st.session_state.welcome_step

    # --- STEP 1: Overview ---
    if step == 1:
        st.markdown("### What This App Does")
        st.markdown(
            "This tool automates the **quarterly actuals compliance review** "
            "for New Mexico school districts and charter schools. It runs "
            "**55+ automated checks** against your uploaded financial reports "
            "and produces a detailed compliance checklist, flagged findings, "
            "and exportable reports (Word memo, Excel tracker, and HTML dashboard)."
        )
        st.markdown(
            "**You'll need three files** to run a complete review. "
            "The next steps will show you where to find each one."
        )
        st.info(
            "💡 **Tip:** You can still use the app with just a Cash Report — "
            "revenue and expenditure checks will simply be skipped."
        )
        col_l, col_r = st.columns(2)
        with col_r:
            if st.button("Next →", use_container_width=True, key="w_next_1"):
                st.session_state.welcome_step = 2
                st.rerun()

    # --- STEP 2: Where to get the files ---
    elif step == 2:
        st.markdown("### Where to Get Your Reports")

        st.markdown("**① Cash Report** (Excel)")
        st.markdown(
            "Download from the district/charter's quarterly submission. "
            "The app reads the **Summary** tab automatically. "
            "Most analysts already have this file from the submission package."
        )

        st.markdown("**② Revenue Report & ③ Expenditure Report** (CSV or Excel)")
        st.markdown(
            "Pull these from the **OBMS Financial Explorer**:"
        )
        st.markdown(
            "👉 [Open OBMS Financial Explorer]"
            "(https://huggingface.co/spaces/bobthehermit/OBMS-Financial-Explorer)"
        )
        st.markdown(
            "In the Explorer app:\n"
            "- Go to the **Actuals** tab\n"
            "- Select the entity and fiscal year\n"
            "- Download the **Revenue Actuals** report (CSV)\n"
            "- Download the **Expenditure Actuals** report (CSV)\n"
        )
        st.warning(
            "⚠️ Make sure your Revenue and Expenditure reports are for the "
            "**same entity and period** as your Cash Report."
        )

        col_l, col_r = st.columns(2)
        with col_l:
            if st.button("← Back", use_container_width=True, key="w_back_2"):
                st.session_state.welcome_step = 1
                st.rerun()
        with col_r:
            if st.button("Next →", use_container_width=True, key="w_next_2"):
                st.session_state.welcome_step = 3
                st.rerun()

    # --- STEP 3: Workflow & saving ---
    elif step == 3:
        st.markdown("### How the Review Works")

        st.markdown(
            "**1. Upload** your three reports in the sidebar.\n\n"
            "**2. Review** — The app runs automated checks and displays "
            "results inside each checklist step. Steps with issues show "
            "🚩 flags; clean steps show ✅.\n\n"
            "**3. Add notes** — Use the notes field in each step to "
            "record your analyst comments, questions for the district, "
            "or follow-up items.\n\n"
            "**4. Export** — Download your finished review as a Word memo, "
            "Excel checklist tracker, or a full HTML visual dashboard."
        )

        st.markdown("---")
        st.markdown("**💾 Saving Your Progress**")
        st.markdown(
            "Your work isn't lost if you close the browser. Use "
            "**\"Save Your Progress\"** in the sidebar to download a session "
            "file. To pick up where you left off, use **\"Resume a Previous "
            "Review\"** and upload that file. It restores your uploaded data, "
            "checklist progress, and all notes."
        )

        col_l, col_r = st.columns(2)
        with col_l:
            if st.button("← Back", use_container_width=True, key="w_back_3"):
                st.session_state.welcome_step = 2
                st.rerun()
        with col_r:
            if st.button("Get Started ✓", type="primary", use_container_width=True, key="w_done"):
                st.session_state.welcome_dismissed = True
                st.session_state.welcome_step = 1  # Reset for next time
                st.rerun()

# ---------- MAIN APP ----------

def main():
    if not st.session_state.welcome_dismissed:
        render_welcome_modal()
    render_header(HEADER_TITLE, HEADER_SUB, LOGO_LEFT_PATH, LOGO_RIGHT_PATH, LOGO_LEFT_LINK, LOGO_RIGHT_LINK, SHOW_HEADER_LOGOS)
    render_sidebar_logo(SIDEBAR_LOGO_PATH)

    # STICKY HEADER CSS
    st.markdown(
        """
        <style>
            div[data-testid="stVerticalBlock"] > div:has(div[data-testid="stProgress"]) {
                position: sticky;
                top: 2.875rem;
                background-color: white;
                z-index: 999;
                padding-top: 10px;
                padding-bottom: 10px;
                border-bottom: 1px solid #f0f2f6;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    with st.sidebar:
        # Help button
        if st.button("❓ How to Use This App", use_container_width=True):
            st.session_state.welcome_dismissed = False
            st.rerun()

        st.divider()

        st.header("1. Review Settings")
        is_q1 = st.radio("Review Period", ["Q1", "Q2-Q4"], index=1) == "Q1"

        st.divider()

        st.header("2. Upload Reports")

        st.markdown("**① Cash Report**")
        st.caption("Excel file from the district's quarterly submission (Summary tab).")
        cash_file = st.file_uploader("Cash Report", type=['csv', 'xlsx', 'xls'],
                                     key='cash', label_visibility="collapsed")
        st.markdown("---")

        st.markdown("**② Revenue Actuals Report**")
        st.caption(
            "CSV/Excel from "
            "[OBMS Financial Explorer](https://huggingface.co/spaces/bobthehermit/OBMS-Financial-Explorer) "
            "→ Actuals tab → Revenue download."
        )
        rev_file = st.file_uploader("Revenue Report", type=['csv', 'xlsx'],
                                    key='rev', label_visibility="collapsed")
        st.markdown("---")

        st.markdown("**③ Expenditure Actuals Report**")
        st.caption(
            "CSV/Excel from "
            "[OBMS Financial Explorer](https://huggingface.co/spaces/bobthehermit/OBMS-Financial-Explorer) "
            "→ Actuals tab → Expenditure download."
        )
        exp_file = st.file_uploader("Expenditure Report", type=['csv', 'xlsx'],
                                    key='exp', label_visibility="collapsed")

        if cash_file:
            st.session_state.cash_df = load_cash_from_excel(cash_file)
        if rev_file:
            st.session_state.revenue_df = load_report_file(rev_file, "Rev")
        if exp_file:
            st.session_state.expenditure_df = load_report_file(exp_file, "Exp")

        st.divider()

        if st.session_state.revenue_df is not None or st.session_state.expenditure_df is not None:
            if not st.session_state.entity_name:
                st.session_state.entity_name = detect_entity_name(
                    st.session_state.revenue_df, st.session_state.expenditure_df
                )
        st.session_state.entity_name = st.text_input("Entity Name", value=st.session_state.entity_name)

        st.divider()

        st.header("3. Save / Resume Progress")
        st.caption(
            "Save your current review (uploaded data, checklist, and notes) "
            "so you can close the browser and pick up later."
        )

        current_state = {
            'checklist_data': st.session_state.checklist_data,
            'notes_by_step': st.session_state.notes_by_step,
            'cash_df': st.session_state.cash_df,
            'revenue_df': st.session_state.revenue_df,
            'expenditure_df': st.session_state.expenditure_df,
            'entity_name': st.session_state.entity_name,
            'step_6_period': st.session_state.get('step_6_period', 0.0),
            'step_6_ytd': st.session_state.get('step_6_ytd', 0.0),
            'step_7_budget': st.session_state.get('step_7_budget', 0.0),
            'step_47_last_year': st.session_state.get('step_47_last_year', 0.0),
            'step_47_this_year': st.session_state.get('step_47_this_year', 0.0),
        }

        buffer = BytesIO()
        pickle.dump(current_state, buffer)
        buffer.seek(0)

        entity_slug = (st.session_state.entity_name.replace(' ', '_')
                       if st.session_state.entity_name else "Review")
        st.download_button(
            label="💾 Save Your Progress",
            data=buffer,
            file_name=f"Review_{entity_slug}_{datetime.now().strftime('%Y%m%d')}.pkl",
            mime="application/octet-stream",
            use_container_width=True,
            help="Downloads a file that stores your entire review session."
        )

        st.markdown("")
        st.markdown("**Resume a Previous Review**")
        st.caption("Upload a previously saved progress file to continue where you left off.")
        uploaded_session = st.file_uploader("Resume session", type=["pkl"],
                                           label_visibility="collapsed")

        if uploaded_session is not None:
            if ('last_loaded_file' not in st.session_state
                    or st.session_state.last_loaded_file != uploaded_session.name):
                try:
                    data = pickle.load(uploaded_session)
                    st.session_state.checklist_data = data.get('checklist_data', [])
                    st.session_state.notes_by_step = data.get('notes_by_step', {})
                    st.session_state.cash_df = data.get('cash_df')
                    st.session_state.revenue_df = data.get('revenue_df')
                    st.session_state.expenditure_df = data.get('expenditure_df')
                    st.session_state.entity_name = data.get('entity_name', "")
                    st.session_state['step_6_period'] = data.get('step_6_period', 0.0)
                    st.session_state['step_6_ytd'] = data.get('step_6_ytd', 0.0)
                    st.session_state['step_7_budget'] = data.get('step_7_budget', 0.0)
                    st.session_state['step_47_last_year'] = data.get('step_47_last_year', 0.0)
                    st.session_state['step_47_this_year'] = data.get('step_47_this_year', 0.0)
                    st.session_state.last_loaded_file = uploaded_session.name
                    st.success("✅ Session restored! Your data and notes are loaded.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error loading session: {e}")

    # --- MAIN CONTENT ---
    if st.session_state.cash_df is not None:
        
        user_inputs = {
            'step_6_period': st.session_state.get('step_6_period', 0.0),
            'step_6_ytd': st.session_state.get('step_6_ytd', 0.0),
            'step_7_budget': st.session_state.get('step_7_budget', 0.0),
            'step_47_last_year': st.session_state.get('step_47_last_year', 0.0),
            'step_47_this_year': st.session_state.get('step_47_this_year', 0.0),
        }

        # Run Validations
        st.session_state.validation_results, st.session_state.table_findings = run_all_validations(
            st.session_state.cash_df,
            st.session_state.revenue_df,
            st.session_state.expenditure_df,
            st.session_state.entity_name,
            is_q1,
            user_inputs
        )
        
        # Generate Analysis Summary
        analysis_summary = generate_analysis_summary(
            st.session_state.cash_df,
            st.session_state.revenue_df,
            st.session_state.expenditure_df,
            st.session_state.entity_name,
            is_q1,
            st.session_state.validation_results
        )

        # Sticky Progress Bar
        total = len(st.session_state.checklist_data)
        done = sum(1 for i in st.session_state.checklist_data if i['completed'])
        st.progress(done / total if total > 0 else 0)
        st.caption(f"Progress: {done}/{total}")
        
        # --- ANALYSIS DASHBOARD ---
        st.header("📊 Analysis Dashboard")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if 'total_revenue' in analysis_summary.get('statistics', {}):
                st.metric("Total Revenue YTD", f"${analysis_summary['statistics']['total_revenue']:,.2f}")
        with col2:
            if 'total_expenditure' in analysis_summary.get('statistics', {}):
                st.metric("Total Expenditure YTD", f"${analysis_summary['statistics']['total_expenditure']:,.2f}")
        with col3:
            if 'total_cash' in analysis_summary.get('statistics', {}):
                st.metric("Total Cash Balance", f"${analysis_summary['statistics']['total_cash']:,.2f}")
        
        # Quick Concerns Summary
        if analysis_summary.get('concerns'):
            with st.expander(f"⚠️ Key Concerns ({len(analysis_summary['concerns'])})", expanded=True):
                for concern in analysis_summary['concerns'][:10]:
                    st.write(f"• {concern}")

        # Filters
        c1, c2 = st.columns([3, 1])
        search = c1.text_input("🔍 Search", "")
        show_incomplete = c2.checkbox("Incomplete Only")

        # Exports
        st.divider()
        ec1, ec2, ec3 = st.columns(3)
        
        memo_bytes = export_findings_memo(
            st.session_state.checklist_data, 
            st.session_state.entity_name, 
            analysis_summary,
            st.session_state.table_findings
        )
        tracker_bytes = export_checklist_tracker(st.session_state.checklist_data)
        
        ec1.download_button("📄 Download Findings Memo (Word)", data=memo_bytes, file_name=f"Memo_{st.session_state.entity_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        ec2.download_button("📊 Download Checklist (Excel)", data=tracker_bytes, file_name=f"Tracker_{st.session_state.entity_name}.xlsx", mime="application/vnd.openxmlformats-officedocument.ms-excel", use_container_width=True)
        # Generate HTML Report
        html_report = generate_html_report(
            entity_name=st.session_state.entity_name,
            revenue_df=st.session_state.revenue_df,
            expenditure_df=st.session_state.expenditure_df,
            cash_df=st.session_state.cash_df,
            validation_results=st.session_state.validation_results,
            table_findings=st.session_state.table_findings,
            notes_by_step=st.session_state.notes_by_step,
            checklist_data=st.session_state.checklist_data
        )
        html_bytes = BytesIO(html_report.encode('utf-8'))
        html_bytes.seek(0)

        ec3.download_button(
            "📊 Download Visual Report (HTML)",
            data=html_bytes,
            file_name=f"Analysis_{st.session_state.entity_name}.html",
            mime="text/html",
            use_container_width=True
        )
        st.divider()

        # Checklist
        items = st.session_state.checklist_data
        if search: items = [i for i in items if search.lower() in i['check'].lower()]
        if show_incomplete: items = [i for i in items if not i['completed']]
        
        areas = list(dict.fromkeys(i['review_area'] for i in items))

        for area in areas:
            st.subheader(f"📂 {area}")
            area_items = [i for i in items if i['review_area'] == area]
            
            for item in area_items:
                step_id = item['step']
                with st.expander(f"Step {step_id}: {item['check']}", expanded=False):
                    
                    # INJECT INPUTS
                    if step_id == 6:
                        st.markdown(f"🔗 [Open June SEG File]({SEG_EXTERNAL_LINK})") 
                        c1, c2 = st.columns(2)
                        st.session_state['step_6_period'] = c1.number_input("SEG Actual Period", value=user_inputs['step_6_period'], key="s6_p")
                        st.session_state['step_6_ytd'] = c2.number_input("SEG Actual YTD", value=user_inputs['step_6_ytd'], key="s6_y")

                    elif step_id == 7:
                        st.session_state['step_7_budget'] = st.number_input("SEG Budgeted", value=user_inputs['step_7_budget'], key="s7_b")
                    
                    elif step_id == 47:
                        c1, c2 = st.columns(2)
                        st.session_state['step_47_last_year'] = c1.number_input("Cash Balance Last Year", value=user_inputs['step_47_last_year'], key="s47_l")
                        st.session_state['step_47_this_year'] = c2.number_input("Cash Balance This Year", value=user_inputs['step_47_this_year'], key="s47_t")

                    # INJECT FINDINGS (Now with PASS/FLAG distinction)
                    findings = st.session_state.validation_results.get(step_id, [])
                    if findings:
                        passes = [f for f in findings if f[0] == "PASS"]
                        flags = [f for f in findings if f[0] == "FLAG"]
                        
                        if flags:
                            st.error("🤖 **Automated Findings (Issues):**")
                            for _, msg in flags:
                                st.write(f"• {msg}")
                        
                        if passes:
                            st.success("🤖 **Automated Findings (Passed):**")
                            for _, msg in passes:
                                st.write(f"• {msg}")
                    
                    # RENDER TABLE FINDINGS
                    render_findings_table(step_id, st.session_state.table_findings)
                    
                    # DISPLAY INFO — compact layout
                    info_col, status_col = st.columns([4, 1])
                    with info_col:
                        # Main check description
                        check_text = inject_links(item['check'])
                        st.markdown(check_text)
                        
                        # Compact metadata line (UCOA + Applies To on one line)
                        ucoa = item['ucoa_line'].strip()
                        applies = item['applies_to'].strip()
                        meta_parts = []
                        if ucoa:
                            meta_parts.append(f"**UCOA:** {ucoa}")
                        if applies:
                            meta_parts.append(f"**Scope:** {applies}")
                        if meta_parts:
                            st.caption(" · ".join(meta_parts))
                        
                        # Support & Method/Notes — show in single clean box when present
                        support_text = item.get('support', '')
                        method_text = item.get('method_notes', '')
                        detail_parts = []
                        if support_text and support_text.strip().lower() not in ('nan', ''):
                            detail_parts.append(inject_links(support_text.strip()))
                        if method_text and method_text.strip().lower() not in ('nan', ''):
                            detail_parts.append(inject_links(method_text.strip()))
                        if detail_parts:
                            st.info("\n\n".join(detail_parts))

                    with status_col:
                        is_done = st.checkbox("Completed", value=item['completed'], key=f"c_{step_id}")
                        idx = next(i for i, x in enumerate(st.session_state.checklist_data) if x['step'] == step_id)
                        st.session_state.checklist_data[idx]['completed'] = is_done
                    
                    # FIXED: Notes with callback for reliable persistence
                    note_key = f"n_{step_id}"
                    current_notes = st.session_state.notes_by_step.get(step_id, "")
                    
                    notes = st.text_area(
                        "Notes", 
                        value=current_notes, 
                        key=note_key, 
                        height=150,  # Increased default height
                        on_change=save_note_callback,
                        args=(step_id,)
                    )
                    
                    # Also save on render (backup)
                    st.session_state.notes_by_step[step_id] = notes
                    st.session_state.checklist_data[idx]['user_notes'] = notes
    else:
        st.markdown("### 👋 Ready to Begin")
        st.markdown(
            "Upload your **Cash Report** in the sidebar to start the review. "
            "For the full suite of automated checks, also upload the "
            "**Revenue** and **Expenditure** Actuals reports."
        )
        st.info(
            "**Need the Revenue or Expenditure reports?** "
            "Download them from the "
            "[OBMS Financial Explorer]"
            "(https://huggingface.co/spaces/bobthehermit/OBMS-Financial-Explorer) "
            "→ Actuals tab."
        )

if __name__ == "__main__":
    main()